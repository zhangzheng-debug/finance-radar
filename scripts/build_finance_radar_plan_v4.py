#!/usr/bin/env python3
"""Build the human-readable Finance Radar V4.0 proposal."""

from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

from build_finance_radar_plan_v3 import (
    BLUE,
    DARK_BLUE,
    GOLD,
    GREEN,
    INK,
    LIGHT,
    LIGHT_BLUE,
    LIGHT_GOLD,
    LIGHT_GREEN,
    LIGHT_RED,
    MUTED,
    NAVY,
    RED,
    add_callout,
    add_figure,
    add_heading,
    add_hyperlink,
    add_list,
    add_num_defs,
    add_page_number,
    add_para,
    add_reference,
    add_table as _base_add_table,
    configure_styles,
    rounded_rectangle,
    set_paragraph_spacing,
    set_repeat_table_header,
    set_run_font,
    set_table_geometry,
    shade_cell,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "financial_event_radar_project_proposal_v4_0_human.docx"
ASSET_DIR = ROOT / "tmp" / "plan_v4_assets"
DEPLOYMENT_IMAGE = ASSET_DIR / "finance_radar_v4_deployment.png"
DEMO_IMAGE = ASSET_DIR / "finance_radar_v4_demo_modes.png"


def add_table(doc, headers, rows, widths_dxa, **kwargs):
    """Add a fixed-width table and keep compact tables on one page when possible."""
    table = _base_add_table(doc, headers, rows, widths_dxa, **kwargs)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        cant_split.set(qn("w:val"), "true")
        tr_pr.append(cant_split)
    if len(table.rows) <= 8:
        for row in table.rows[:-1]:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.keep_with_next = True
    return table


def configure_page_v4(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    hp = section.header.paragraphs[0]
    set_paragraph_spacing(hp, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.LEFT)
    run = hp.add_run("FINANCE RADAR  ·  WEB SITUATION ROOM & EVIDENCE AGENT  ·  V4.0")
    set_run_font(run, size=8.5, bold=True, color=MUTED)

    fp = section.footer.paragraphs[0]
    set_paragraph_spacing(fp, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    run = fp.add_run("北京林业大学理学院2026实训  ·  研究监控用途 / 无交易权限  ·  ")
    set_run_font(run, size=9, color=MUTED)
    add_page_number(fp)


def _fonts():
    normal = r"C:\Windows\Fonts\msyh.ttc"
    bold = r"C:\Windows\Fonts\msyhbd.ttc"
    return {
        "title": ImageFont.truetype(bold, 48),
        "section": ImageFont.truetype(bold, 29),
        "body": ImageFont.truetype(normal, 23),
        "small": ImageFont.truetype(normal, 19),
        "tiny": ImageFont.truetype(normal, 17),
    }


def _center_text(draw, box, text, font, color):
    x1, y1, x2, y2 = box
    bounds = draw.multiline_textbbox((0, 0), text, font=font, spacing=5, align="center")
    width = bounds[2] - bounds[0]
    height = bounds[3] - bounds[1]
    draw.multiline_text(
        (x1 + (x2 - x1 - width) / 2, y1 + (y2 - y1 - height) / 2),
        text,
        font=font,
        fill=color,
        spacing=5,
        align="center",
    )


def _arrow(draw, start, end, color="#2E74B5", width=7):
    x1, y1 = start
    x2, y2 = end
    draw.line((x1, y1, x2, y2), fill=color, width=width)
    if abs(x2 - x1) >= abs(y2 - y1):
        sign = 1 if x2 > x1 else -1
        draw.polygon([(x2, y2), (x2 - 18 * sign, y2 - 12), (x2 - 18 * sign, y2 + 12)], fill=color)
    else:
        sign = 1 if y2 > y1 else -1
        draw.polygon([(x2, y2), (x2 - 12, y2 - 18 * sign), (x2 + 12, y2 - 18 * sign)], fill=color)


def make_deployment_diagram(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fonts = _fonts()
    image = Image.new("RGB", (1800, 1120), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    draw.text((60, 42), "Finance Radar V4.0：新加坡服务器部署拓扑", font=fonts["title"], fill="#17365D")
    draw.text((62, 107), "一个仓库、一个镜像、一个数据契约；按职责运行 Web / API / Worker", font=fonts["body"], fill="#667085")

    rounded_rectangle(draw, (55, 220, 330, 780), 24, "#EAF2F8", "#2E74B5", 4)
    _center_text(draw, (75, 240, 310, 310), "外部只读来源", fonts["section"], "#1F4D78")
    source_lines = ["SEC / EDGAR", "监管与公司公告", "OpenNews / RSS", "Telegram线索", "Binance / IBKR行情"]
    y = 350
    for line in source_lines:
        draw.text((86, y), "• " + line, font=fonts["small"], fill="#1E293B")
        y += 72

    rounded_rectangle(draw, (420, 160, 1460, 920), 28, "#F8FAFC", "#17365D", 5)
    draw.text((470, 190), "新加坡 VPS · Docker Compose", font=fonts["section"], fill="#17365D")

    boxes = [
        ((480, 290, 760, 435), "Caddy HTTPS", "证书 / 反向代理\n访问控制", "#EAF2F8", "#2E74B5"),
        ((820, 290, 1110, 435), "Web 情报台", "Streamlit 四页\n3–5秒刷新", "#EDF7F0", "#217346"),
        ((1170, 290, 1400, 435), "FastAPI", "只读查询\nSchema约束", "#EAF2F8", "#2E74B5"),
        ((480, 520, 785, 710), "后台 Worker", "持续采集 / 去重\n证据链 / Outbox", "#FFF7E6", "#9A6700"),
        ((850, 520, 1115, 710), "小模型", "风险路由\nShadow + ABSTAIN", "#FFF7E6", "#9A6700"),
        ((1180, 520, 1400, 710), "Telegram", "高优先级通知\n网页深链接", "#EDF7F0", "#217346"),
        ((480, 770, 785, 875), "SQLite WAL", "正式历史账本 / 单写者", "#F4F6F9", "#667085"),
        ((850, 770, 1115, 875), "证据快照", "内容哈希 / 原始文件", "#F4F6F9", "#667085"),
        ((1180, 770, 1400, 875), "备份", "日备份 + SHA256 + 恢复", "#F4F6F9", "#667085"),
    ]
    for box, title, body, fill, outline in boxes:
        rounded_rectangle(draw, box, 20, fill, outline, 3)
        x1, y1, x2, y2 = box
        _center_text(draw, (x1 + 8, y1 + 10, x2 - 8, y1 + 66), title, fonts["section"], outline)
        _center_text(draw, (x1 + 10, y1 + 66, x2 - 10, y2 - 8), body, fonts["small"], "#1E293B")

    rounded_rectangle(draw, (1530, 250, 1750, 470), 24, "#EDF7F0", "#217346", 4)
    _center_text(draw, (1545, 268, 1735, 450), "浏览器\n主展示终端", fonts["section"], "#217346")
    rounded_rectangle(draw, (1530, 560, 1750, 755), 24, "#EAF2F8", "#2E74B5", 4)
    _center_text(draw, (1545, 578, 1735, 735), "本地 Windows\n离线 Replay 兜底", fonts["section"], "#1F4D78")

    _arrow(draw, (330, 500), (470, 500))
    _arrow(draw, (760, 362), (820, 362))
    _arrow(draw, (1110, 362), (1170, 362))
    _arrow(draw, (1460, 365), (1530, 365))
    _arrow(draw, (630, 500), (630, 520), "#9A6700")
    _arrow(draw, (785, 615), (850, 615), "#9A6700")
    _arrow(draw, (1115, 615), (1180, 615), "#217346")
    _arrow(draw, (630, 710), (630, 770), "#667085")
    _arrow(draw, (980, 710), (980, 770), "#667085")
    _arrow(draw, (1290, 710), (1290, 770), "#667085")
    _arrow(draw, (1530, 655), (1460, 655), "#2E74B5")

    draw.text((60, 1010), "核心原则：网页不直连数据库；Worker唯一写入正式账本；Replay使用隔离账本；任何组件均无交易接口。", font=fonts["body"], fill="#667085")
    image.save(path, quality=95)


def make_demo_diagram(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fonts = _fonts()
    image = Image.new("RGB", (1800, 760), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    draw.text((60, 40), "答辩不押注现场新闻：三模式演示协议", font=fonts["title"], fill="#17365D")
    draw.text((62, 105), "三个模式证明不同问题，页面必须始终显示当前模式与数据时间", font=fonts["body"], fill="#667085")

    cards = [
        ((75, 230, 535, 610), "LIVE", "证明外部链路活着", ["来源健康与心跳", "最近一次成功采集", "新事件出现则即时展示"], "#EDF7F0", "#217346"),
        ((670, 230, 1130, 610), "RECENT CAPTURE", "证明真实数据不是伪造", ["展示近24–72小时捕获", "保留原文、时间和哈希", "解释抓取与入账轨迹"], "#EAF2F8", "#2E74B5"),
        ((1265, 230, 1725, 610), "REPLAY", "证明完整能力可复现", ["NINEQ正例 + WOLF2反例", "同一套下游代码", "独立账本且醒目标识"], "#FFF7E6", "#9A6700"),
    ]
    for box, title, subtitle, lines, fill, outline in cards:
        rounded_rectangle(draw, box, 28, fill, outline, 4)
        x1, y1, x2, _ = box
        _center_text(draw, (x1 + 10, y1 + 20, x2 - 10, y1 + 90), title, fonts["section"], outline)
        _center_text(draw, (x1 + 20, y1 + 100, x2 - 20, y1 + 160), subtitle, fonts["body"], "#1E293B")
        y = y1 + 205
        for line in lines:
            draw.text((x1 + 42, y), "• " + line, font=fonts["small"], fill="#1E293B")
            y += 58
    _arrow(draw, (535, 420), (670, 420), "#667085")
    _arrow(draw, (1130, 420), (1265, 420), "#667085")
    draw.text((60, 680), "结论：LIVE 不承担“恰好有 SEC 大事件”的赌注；REPLAY 是主演示，LIVE 与 RECENT 是真实性证据。", font=fonts["body"], fill="#667085")
    image.save(path, quality=95)


def add_cover(doc: Document) -> None:
    add_para(doc, "HIGH-DIFFICULTY PROJECT PROPOSAL  ·  V4.0", align=WD_ALIGN_PARAGRAPH.CENTER,
             before=34, after=16, line=1.0, size=10.5, color=BLUE)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=9, line=1.05, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run("基于多源证据链与实时 Web 控制台的")
    set_run_font(run, size=21, bold=True, color=NAVY)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=10, line=1.05, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run("金融事件情报 Agent")
    set_run_font(run, size=30, bold=True, color=NAVY)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=8, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run("全极性事实核验 · 重大下行风险路由 · 历史证据账本 · 可审计回放")
    set_run_font(run, size=13, color=DARK_BLUE)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=24, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run("自主高难度创新选题申报稿 / 95+ 交付蓝图")
    set_run_font(run, size=11, italic=True, color=MUTED)

    table = doc.add_table(rows=3, cols=2)
    values = [
        ("版本", "V4.0", "编制日期", "2026年7月17日"),
        ("主展示", "Web Situation Room", "通知出口", "Telegram 深链接"),
        ("部署目标", "新加坡 VPS + 本地回放", "交易权限", "无；全部只读"),
    ]
    for row_idx, values_row in enumerate(values):
        for col_idx in range(2):
            label = values_row[col_idx * 2]
            value = values_row[col_idx * 2 + 1]
            cell = table.rows[row_idx].cells[col_idx]
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, before=2, after=2, line=1.1, align=WD_ALIGN_PARAGRAPH.CENTER)
            r1 = p.add_run(f"{label}\n")
            set_run_font(r1, size=8.5, bold=True, color=MUTED)
            r2 = p.add_run(value)
            set_run_font(r2, size=10.2, bold=True, color=NAVY)
            shade_cell(cell, LIGHT)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [4680, 4680])

    add_para(doc, "", after=18, line=1.0)
    add_callout(
        doc,
        "一句话价值",
        "系统不预测涨跌，也不执行交易；它持续发现金融事件，保存原始历史，拆解原子声明，寻找一手证据，解释支持、反驳与不足，并把完整轨迹呈现在网页情报台。",
        fill=LIGHT_BLUE,
    )
    add_para(
        doc,
        "V4.0 严格区分“当前已验证”和“计划交付”。服务器、网页、小模型和回放只有在实际部署、测试和留痕后才可作为成果申报。",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=14,
        after=0,
        line=1.0,
        size=9.5,
        color=MUTED,
    )
    doc.add_page_break()


def build() -> Path:
    make_deployment_diagram(DEPLOYMENT_IMAGE)
    make_demo_diagram(DEMO_IMAGE)
    doc = Document()
    configure_page_v4(doc)
    configure_styles(doc)
    bullet_id, decimal_id = add_num_defs(doc)
    props = doc.core_properties
    props.title = "基于多源证据链与实时Web控制台的金融事件情报Agent"
    props.subject = "北京林业大学理学院2026实训自主高难度创新项目任务书V4.0"
    props.author = "Finance Radar Project Team"
    props.keywords = "Finance Radar, Web Situation Room, Evidence Agent, Replay, Audit, No Trading"

    add_cover(doc)

    add_heading(doc, "执行摘要", 1)
    add_para(doc, "Finance Radar 的核心不是“聚合更多新闻”，而是把消息变成可复核的事件版本：谁在什么时间发布了什么，哪些原子声明得到哪份一手材料支持，哪些仍有冲突或证据不足，证券身份和法律阶段是否发生变化。V4.0 在既有事件账本与安全规则之上，增加网页版情报终端、服务器持续运行、历史持久化、备份恢复、小模型风险路由和三模式可审计演示，使项目从研究原型升级为可部署课程产品。")
    add_callout(doc, "申报判断", "题目本身具备自主高难度创新级潜力；当前成品仍不是95分交付态。完成本任务书的硬门后，92–96分是现实竞争区间；要稳定争取95+，还必须在Day 1获得高难度认定，并通过过程规范、禁飞区讲解、Bug盲测、即兴修改和个人代码理解考核。", fill=LIGHT_GREEN, accent=GREEN)

    add_heading(doc, "V4.0 的五个不可替代升级", 2)
    for item in [
        "主终端升级：Web Situation Room 成为正式展示面，Telegram 仅负责高优先级提醒和深链接。",
        "运行形态升级：新加坡 VPS 7×24小时运行 HTTPS、Web、只读 API、采集 Worker 与模型。",
        "数据生命周期升级：服务器持久卷保存历史账本、证据快照、模型轨迹，具有日备份、哈希和恢复演练。",
        "AI闭环升级：LLM Evidence Agent 负责声明与证据语义工作；CPU小模型只做重大下行风险审核路由。",
        "答辩确定性升级：LIVE、RECENT_CAPTURE、REPLAY 分工，彻底消除“现场恰好没有SEC事件”的演示风险。",
    ]:
        add_list(doc, item, bullet_id)

    add_heading(doc, "阅读导航", 2)
    add_para(doc, "产品与网页见第2–3章；服务器、账本与恢复见第4–5章；模型与AI边界见第6–7章；演示协议见第8章；禁飞区、工程过程和95+验收见第10–17章。", size=9.8, color=MUTED, align=WD_ALIGN_PARAGRAPH.LEFT)

    add_heading(doc, "1. 当前基线与V4目标", 1)
    add_heading(doc, "1.1 已验证的真实基础", 2)
    add_table(
        doc,
        ["能力 / 证据", "当前核验值", "状态"],
        [
            ("SQLite 数据模式", "Schema 12；原始观测、修订、事件版本、证据、任务、outbox、租约、来源游标", "已实现"),
            ("来源与数据", "18个来源；1,160个规范事件；3,556条原始观测；710条来源修订", "已实现"),
            ("审计历史", "2,081个事件版本；2,386条证据；1,898条只读行情审计指标", "已实现"),
            ("自动测试", "232项测试全过；安全审计PASS；交易违规0", "已实现"),
            ("VPS/Web/API", "尚无持续Worker、Web终端、FastAPI、Compose运行证据", "计划实现"),
            ("运行时模型", "尚无实际LLM闭环和训练后小模型；792条裁决为下行风险语料，不是交易标签", "计划实现"),
            ("课程过程", "有效Git、.agent、coverage和正式Sprint证据尚未建立", "待补齐"),
        ],
        [2700, 5060, 1600],
        font_size=8.8,
        center_cols=(2,),
        status_col=2,
    )
    add_callout(doc, "真实性口径", "“可调用过API”只证明外部依赖可用；“在服务器持续运行、失败可恢复、结果可追溯”才算工程成果。本文所有未落地项都标为计划，不把架构图冒充成现状。", fill=LIGHT_GOLD, accent=GOLD)

    add_heading(doc, "1.2 最小成功与95+成功", 2)
    add_table(
        doc,
        ["层级", "必须看到的结果", "评分含义"],
        [
            ("最小可交付", "一条真实事件从采集、入账、证据核验到网页呈现可复现；Telegram可深链", "完成产品闭环"),
            ("优秀交付", "VPS稳定、历史可恢复、双案例Replay、小模型有指标、核心覆盖率≥80%", "具备90–95竞争力"),
            ("95+交付", "高难度审批 + 三禁飞区掌握 + 过程证据 + Bug/即兴修改 + 强答辩", "具备95+竞争力"),
        ],
        [1900, 4960, 2500],
        font_size=9.1,
    )

    add_heading(doc, "2. 产品定位与用户场景", 1)
    add_heading(doc, "2.1 产品定义", 2)
    add_para(doc, "Finance Radar 是面向研究者的金融事件情报工作台。它服务于“快速发现—谨慎核验—保存历史—解释结论—复盘演示”五个动作，不是新闻门户、聊天机器人、行情交易终端或投资建议系统。系统覆盖正面、负面和中性事件；只有小模型训练目标聚焦重大下行风险，因为此类事件通常更紧迫、损失不对称、需要优先审核。")
    add_heading(doc, "2.2 四个核心场景", 2)
    add_table(
        doc,
        ["场景", "用户任务", "系统响应"],
        [
            ("实时监控", "查看来源是否健康、最近发生了什么", "3–5秒刷新事件流，显示模式、时间、新鲜度和来源健康"),
            ("事件核验", "判断说法是否有一手证据", "原子声明—证据矩阵；支持、反驳、证据不足分别展示"),
            ("历史复盘", "理解事件如何修订、证券身份如何变化", "版本时间线、原始快照、哈希、人工覆盖与trace_id"),
            ("演示与教学", "在无新事件时稳定展示完整链路", "冻结真实案例进入同一下游代码，写入隔离Replay账本"),
        ],
        [1650, 3380, 4330],
        font_size=9.0,
    )
    add_callout(doc, "边界声明", "事实极性、事件严重度、资产影响和模型风险路由是四个不同字段。正面事件可以被采集和核验；负面模型偏置被定义为有意的审核优先级，而不是做空信号。", fill=LIGHT_BLUE)

    add_heading(doc, "3. Web Situation Room：主展示终端", 1)
    add_para(doc, "网页不做华而不实的大屏，而做四个可以讲透、可以操作、可以验收的页面。所有页面通过FastAPI读取经过Schema约束的数据，不直接访问数据库；标题栏持续显示LIVE / RECENT / REPLAY、服务器时间和数据时间。")
    add_table(
        doc,
        ["页面", "主要内容", "答辩价值"],
        [
            ("Operations Overview", "实时事件流、来源健康、Worker心跳、队列深度、最后成功采集、模式标识", "证明系统真在运行，而非静态页面"),
            ("Event Intelligence", "事件状态、极性、严重度、风险路由、声明—证据矩阵、版本时间线、行情审计图", "证明模型结论可追溯、可反驳"),
            ("Replay Lab", "选择NINEQ/WOLF2/近期案例，逐步播放发现、拒绝、确认和更新", "保证主演示稳定，展示正反例"),
            ("Model & System Health", "模型卡、版本哈希、Precision/Recall/F1、ABSTAIN、错误、备份与恢复状态", "证明AI与运维均受治理"),
        ],
        [1800, 4760, 2800],
        font_size=8.9,
    )
    add_heading(doc, "3.1 Telegram 的新职责", 2)
    for item in [
        "只推送高优先级摘要、来源状态异常和人工复核请求。",
        "每张卡片携带event_id、数据时间、证据状态与网页版详情深链接。",
        "不作为唯一终端，不承载完整证据矩阵，不替代历史检索。",
        "继续通过outbox保证幂等；重复发送目标为0。",
    ]:
        add_list(doc, item, bullet_id)

    doc.add_page_break()
    add_heading(doc, "4. 总体架构与服务器部署", 1)
    add_figure(doc, DEPLOYMENT_IMAGE, "图1  Finance Radar V4.0 模块化单体与新加坡VPS部署拓扑")
    add_heading(doc, "4.1 为什么选择模块化单体", 2)
    add_para(doc, "V4.0 使用一个仓库、一个应用镜像和一个数据库契约，按职责启动worker、api、web三个进程，由Caddy提供HTTPS。这比在12天内引入React、Redis、Celery、Kafka或Kubernetes更可控，也更容易逐行解释、测试、恢复和现场修改。专业度来自闭环可靠，不来自技术名词数量。")
    add_table(
        doc,
        ["组件", "职责", "硬边界"],
        [
            ("Caddy", "HTTPS证书、反向代理、基础访问控制", "不处理业务数据"),
            ("Worker", "持续采集、先落盘、候选抽取、证据流程、模型路由、outbox", "正式账本唯一写者"),
            ("FastAPI", "只读查询、健康检查、Replay管理员入口", "禁止任意SQL、URL和提示词"),
            ("Streamlit", "四页网页情报台与可视化", "不直连数据库、不伪装交易屏"),
            ("Risk Router", "输出RISK_REVIEW / NON_TARGET / ABSTAIN", "不判断真假、不输出SHORT"),
            ("Ledger", "原始观测、版本、证据、模型运行、通知历史", "不保存密钥、不连交易账户"),
        ],
        [1700, 4720, 2940],
        font_size=8.8,
    )

    add_heading(doc, "4.2 只读API契约", 2)
    add_table(
        doc,
        ["方法", "路径", "用途"],
        [
            ("GET", "/api/v1/health", "服务、数据库、Worker和模型健康"),
            ("GET", "/api/v1/sources", "来源状态、新鲜度与最近错误"),
            ("GET", "/api/v1/events", "筛选、分页、模式隔离的事件列表"),
            ("GET", "/api/v1/events/{event_id}", "事件详情、声明、证据、时间线和trace"),
            ("GET", "/api/v1/model-card", "模型版本、指标、标签定义和限制"),
            ("POST", "/api/v1/demo/replays/{scenario_id}", "仅管理员触发隔离Replay"),
        ],
        [1100, 4200, 4060],
        font_size=9.0,
        center_cols=(0,),
    )

    add_heading(doc, "5. 历史账本、备份与恢复", 1)
    add_heading(doc, "5.1 12天核心仍采用SQLite WAL", 2)
    add_para(doc, "现有约19.5MB数据库已具备30表左右的成熟账本语义、WAL、任务、租约、游标和outbox。V4.0不为“看起来像大厂”重写存储层；生产核心采用VPS持久卷上的SQLite WAL，由Worker单写，API只读。PostgreSQL作为仓储适配器升级项，只有在并发写者、用户数、锁等待或数据规模超过触发条件时迁移。")
    add_callout(doc, "数据库专业性", "评委真正能验收的是：重启不丢数据、备份一致、哈希可验、十分钟内可恢复、Replay不污染正式历史。数据库品牌本身不加分。", fill=LIGHT_GREEN, accent=GREEN)
    add_table(
        doc,
        ["对象", "保存内容", "保留 / 验收"],
        [
            ("正式账本", "raw observations、revisions、events、versions、evidence、model_runs", "持久卷；重启后游标与outbox不丢"),
            ("证据快照", "原文或文件、内容哈希、抓取时间、URL、MIME", "内容寻址；可回到原始证据"),
            ("Replay账本", "冻结案例的隔离运行结果", "独立数据库；页面永远显示REPLAY"),
            ("日/周备份", "SQLite在线一致性备份 + SHA256清单", "14个日备份、8个周备份"),
            ("恢复演练", "恢复到临时目录并执行审计与样例查询", "每版至少一次；目标10分钟内"),
        ],
        [1850, 4860, 2650],
        font_size=8.9,
    )
    add_heading(doc, "5.2 PostgreSQL升级触发条件", 2)
    for item in [
        "出现两个以上并发写入Worker；",
        "需要十名以上并发用户；",
        "数据库锁等待P95持续高于50ms；",
        "数据达到数十GB或需要跨节点高可用。",
    ]:
        add_list(doc, item, bullet_id)

    add_heading(doc, "6. Evidence Agent：AI进入核心业务", 1)
    add_para(doc, "LLM不是聊天层，而是受Schema、预算、工具白名单和证据门约束的研究节点。它对每条候选消息拆分原子声明，规划需要查询的一手材料，把证据片段映射到声明，并生成带引用摘要。任何格式错误、来源中断、证据冲突、最终性不足或预算耗尽都必须fail closed。")
    add_table(
        doc,
        ["节点", "模型工作", "确定性代码工作"],
        [
            ("Claim Extractor", "拆分主体、动作、时间、数值、阶段和资产", "JSON Schema校验、字段规范化、输入哈希"),
            ("Evidence Planner", "提出需要的权威证据与查询计划", "来源白名单、域名策略、预算和超时"),
            ("Relation Judge", "判断片段支持、反驳或无关", "引用存在性、来源等级、独立性和最终性门"),
            ("Cited Summarizer", "生成逐声明带引用的简报", "禁止无引用结论、版本化、人工覆盖审计"),
        ],
        [1900, 3770, 3690],
        font_size=8.9,
    )
    add_callout(doc, "Fail closed", "模型异常时输出INSUFFICIENT或HUMAN_REVIEW；不能猜测性发布，不能绕过S级限制，不能调用交易工具。", fill=LIGHT_RED, accent=RED)

    add_heading(doc, "6.1 Agent运行时可观测性", 2)
    for item in [
        "每次运行记录trace_id、输入哈希、模型/提示词版本、工具调用、延迟、成本、输出Schema和人工覆盖。",
        "网页事件详情可按时间回放“发现→抽取→检索→关系判断→安全门→发布”。",
        "所有外部工具均只读，URL、查询次数、超时和返回大小受白名单与预算限制。",
    ]:
        add_list(doc, item, bullet_id)

    add_heading(doc, "7. 服务器小模型：重大下行风险审核路由", 1)
    add_heading(doc, "7.1 任务定义，而不是交易模型", 2)
    add_para(doc, "历史792条人工裁决主要围绕破产、退市、重组失败、欺诈、监管处罚等不利事件。V4.0不把它包装成做空模型，而把它定义为“重大下行风险审核路由器”：帮助人工队列优先查看可能存在重大不利影响的事件。全系统仍采集正面、中性和负面事件；正面/中性样本作为NON_TARGET对照，模糊样本进入ABSTAIN。")
    add_table(
        doc,
        ["项目", "V4.0契约"],
        [
            ("模型", "word + char TF-IDF + LogisticRegression(class_weight='balanced')"),
            ("输出", "RISK_REVIEW / NON_TARGET / ABSTAIN + probability + model_version + input_hash"),
            ("正类", "在定义窗口与证据标准下，需要优先人工审核的重大不利事件"),
            ("非目标类", "正面、中性、轻微、已拒绝或不相关事件"),
            ("禁止输出", "LONG、SHORT、目标价、预期收益、自动严重度、发布许可"),
            ("部署", "Worker进程内CPU推理；模型工件、模型卡、特征Schema、SHA256一起发布"),
            ("失效策略", "加载失败、漂移或低置信度一律ABSTAIN；不改变事实核验"),
        ],
        [2100, 7260],
        font_size=9.0,
    )
    add_heading(doc, "7.2 评测与是否上线", 2)
    add_para(doc, "数据按时间、发行人和事件链分组切分，禁止同一事件修订泄漏到训练与测试两侧。与关键词规则、多数类和随机基线比较，报告Precision、Recall、F1、PR-AUC、校准误差和ABSTAIN覆盖。若冻结测试集没有稳定超过规则基线，模型仍可在服务器Shadow Mode运行并记录结果，但不得改变审核顺序。")
    add_callout(doc, "可解释的负面聚焦", "重大下行风险常具有损失不对称与审核时效性，因此优先路由是合理产品目标；这只能解释“审核优先级”，不能解释为“负面事件一定导致价格下跌”。", fill=LIGHT_GOLD, accent=GOLD)

    add_heading(doc, "8. 三模式演示：把外部不确定性变成确定证据", 1)
    add_figure(doc, DEMO_IMAGE, "图2  LIVE、RECENT_CAPTURE与REPLAY三模式答辩协议")
    add_heading(doc, "8.1 主演示顺序（3分钟）", 2)
    add_table(
        doc,
        ["时间", "动作", "必须说清的结论"],
        [
            ("0:00–0:25", "打开Operations Overview，展示HTTPS、服务器时间、来源健康和Worker心跳", "系统当前真实在线；不是录像"),
            ("0:25–0:55", "打开RECENT案例，展示原始快照、抓取时间、哈希和event_id", "真实数据已进入历史账本"),
            ("0:55–2:15", "运行NINEQ Replay，查看声明—证据矩阵、时序身份、最终性和版本更新", "同一套代码完成完整正例链路"),
            ("2:15–2:45", "运行WOLF2反例，展示冲突、拒绝和ABSTAIN", "系统知道何时不下结论"),
            ("2:45–3:00", "展示模型卡、备份状态和Telegram深链接", "AI、工程与通知闭环"),
        ],
        [1350, 4580, 3430],
        font_size=8.9,
    )
    add_callout(doc, "演示原则", "LIVE只证明链路健康，不承担“恰好发生SEC重大事件”的赌注；REPLAY是主演示，案例来自冻结真实数据，并且与正式账本视觉和物理隔离。", fill=LIGHT_GREEN, accent=GREEN)

    add_heading(doc, "9. 数据、事件与安全契约", 1)
    add_heading(doc, "9.1 必须分开的四层语义", 2)
    add_table(
        doc,
        ["字段组", "回答的问题", "示例"],
        [
            ("event_polarity", "事件内容本身偏正面、负面还是中性？", "FAVORABLE / ADVERSE / NEUTRAL / MIXED / UNKNOWN"),
            ("severity/finality", "事件是否重大、是否已达到法律或业务最终阶段？", "申请≠获批≠生效≠注销"),
            ("asset_impact", "可能影响哪些资产与机制？", "只做解释，不把波动当因果证明"),
            ("risk_route", "是否需要优先人工审核？", "RISK_REVIEW / NON_TARGET / ABSTAIN"),
        ],
        [1800, 4250, 3310],
        font_size=9.0,
    )
    add_heading(doc, "9.2 不可逾越的安全边界", 2)
    for item in [
        "仓库、服务器、网页和Telegram均不提供下单、撤单、账户、资金、仓位或API签名能力。",
        "不读取、不挂载、不修改新加坡服务器上既有量化交易目录；加密行情只使用公共只读接口。",
        "密钥只保存在服务器环境变量或secret文件，文档、日志、数据库、截图和Git中不得出现。",
        "行情只用于时间对齐和事后审计，不用于反向证明事件真假或宣传因果收益。",
        "Replay必须明显标识，不得伪装成LIVE；模型异常一律降级，不得静默失败。",
    ]:
        add_list(doc, item, bullet_id)

    add_heading(doc, "10. 三个AI禁飞区与代码理解", 1)
    add_para(doc, "禁飞区必须缩小为纯逻辑内核：无网络、无数据库、无LLM、输入输出稳定、约100–150行，学生手写并逐行解释。Day 1向教师提交一页边界说明并获得书面确认；若自定义第三禁飞区不获认可，则切换到课程标准的多样性重排内核。")
    add_table(
        doc,
        ["禁飞区", "职责", "现场必须解释"],
        [
            ("event_fingerprint.py", "同一事件去重、修订判定、时点身份键", "标准化、哈希、误合并/漏合并边界"),
            ("evidence_gate.py", "支持/反驳/不足、引用完整性、来源独立性", "逐声明聚合、阈值、fail closed"),
            ("finality_gate.py", "阶段状态转换、最终性、禁止模型自动升S", "状态机、不变量、非法转换"),
        ],
        [2300, 3800, 3260],
        font_size=9.0,
    )
    add_callout(doc, "禁飞区不是装饰", "讲不清的模块可能被清零或要求删除重写。外围适配器可由AI协助，但禁飞区必须保留设计草图、手写提交、单元测试和个人讲解记录。", fill=LIGHT_RED, accent=RED)

    add_heading(doc, "11. Bug注入、即兴修改与答辩韧性", 1)
    add_heading(doc, "11.1 内部盲测，不冒充官方题目", 2)
    add_table(
        doc,
        ["类别", "内部练习题", "定位证据"],
        [
            ("基础Bug", "时间戳时区错误导致新鲜度为负；event_id大小写不一致", "失败测试、结构化日志、最小复现"),
            ("边界Bug", "同文修订被误判新事件；证据只有转载却越过独立性门", "fixture、门函数输出、trace"),
            ("集成Bug", "outbox重试重复发送；Replay误写正式账本", "幂等键、事务、数据库差异"),
        ],
        [1700, 4930, 2730],
        font_size=8.9,
    )
    add_para(doc, "每个Sprint由非实现者盲注3个Bug，30分钟内定位并修复；每位成员至少完成3次演练，优秀目标是30分钟修复2–3个并讲清根因、证据、补丁和回归测试。")
    add_heading(doc, "11.2 即兴修改扩展点", 2)
    for item in [
        "事件类型注册表：新增事件类型不改核心状态机；",
        "网页筛选Schema：新增“证据不足”或来源过滤器；",
        "风险阈值配置：修改阈值必须同步展示模型版本与测试；",
        "Replay场景注册：新增冻结案例只加适配器和fixture；",
        "来源适配器接口：新增官方源不修改账本和证据门。",
    ]:
        add_list(doc, item, bullet_id)

    add_heading(doc, "12. 工程规范、测试与可观测性", 1)
    add_table(
        doc,
        ["领域", "V4.0硬门", "验收证据"],
        [
            ("Git", "Day 1从当前baseline建立有效仓库；细粒度提交，>300行说明", "commit、PR/评审、贡献矩阵"),
            (".agent", "保存架构、接口、数据、测试、部署、AI日志与每Sprint反思", "目录与文档完整"),
            ("测试", "现有232项不退化；核心覆盖率≥80%；契约/集成/恢复/安全测试", "pytest、coverage、报告"),
            ("可观测", "结构化日志、trace_id、来源新鲜度、队列、模型延迟、备份状态", "网页健康页与日志样例"),
            ("部署", "Docker Compose一键启动；restart policy；HTTPS；最小权限", "VPS smoke test与重启测试"),
            ("恢复", "备份哈希、临时恢复、审计PASS、十分钟目标", "恢复演练记录"),
        ],
        [1500, 4940, 2920],
        font_size=8.8,
    )
    add_heading(doc, "12.1 五份必须由人主导的设计产物", 2)
    for item in [
        "总体架构与组件边界；",
        "API/事件/错误契约；",
        "数据模型、迁移和备份恢复设计；",
        "测试策略、禁飞区用例和故障注入表；",
        "VPS部署、安全、回滚与演示Runbook。",
    ]:
        add_list(doc, item, decimal_id)

    add_heading(doc, "13. 12天 / 3个Sprint实施计划", 1)
    add_table(
        doc,
        ["时间", "核心交付", "Sprint Gate"],
        [
            ("Day 1", "高难度评审；冻结V4范围与3禁飞区；建立有效Git、.agent、角色与Backlog", "教师确认题级与禁飞区；baseline可复跑"),
            ("Day 2–3", "Schema 13；FastAPI健康/事件契约；VPS普通用户、Caddy、Compose；备份恢复脚本", "公网HTTPS健康页；重启与恢复测试"),
            ("Day 4", "Sprint 1 Review/Retro；盲注3 Bug；接口和数据库设计封版", "过程证据完整；不带红项进入Sprint 2"),
            ("Day 5–6", "Worker常驻化；只读API；Telegram网页深链；Web Overview/Event页", "新事件P95≤5秒可见；数据重启不丢"),
            ("Day 7–8", "Replay Lab；NINEQ/WOLF2；Evidence Agent实际运行；第二轮Bug", "正反例均可回放；trace完整"),
            ("Day 9", "Sprint 2 Review/Retro；核心覆盖率、API契约、恢复演练", "覆盖率≥80%；安全审计PASS"),
            ("Day 10", "小模型训练、模型卡、Shadow部署；Model & Health页", "冻结测试与规则基线对照；异常ABSTAIN"),
            ("Day 11", "24–48小时运行证据；第三轮Bug；每人即兴修改演练", "LIVE/RECENT/REPLAY稳定；Runbook可执行"),
            ("Day 12", "3分钟答辩、禁飞区走查、材料封版、Sprint 3 Review/Retro", "个人讲解与现场修改通过"),
        ],
        [1250, 5120, 2990],
        font_size=8.6,
    )
    add_callout(doc, "顺序纪律", "服务器部署必须在Day 2–3得到第一份运行证据，不能拖到最后。若Web/API/恢复/Replay尚未闭环，暂停PostgreSQL、WebSocket、React、ONNX等增强项。", fill=LIGHT_GOLD, accent=GOLD)

    add_heading(doc, "14. 定量验收指标", 1)
    add_table(
        doc,
        ["验收项", "通过标准", "证据"],
        [
            ("公网访问", "HTTPS网页可从外网打开；访问控制有效", "浏览器实测、TLS信息"),
            ("新鲜度", "新事件落账后网页P95可见延迟≤5秒", "自动时间戳测试"),
            ("查询", "1,000条事件查询P95≤1秒", "基准报告"),
            ("持久化", "VPS/容器重启后事件、游标、outbox不丢", "前后校验"),
            ("恢复", "备份在10分钟内恢复并通过审计", "恢复Runbook记录"),
            ("Replay", "NINEQ和WOLF2各≤30秒；不污染正式账本", "trace与DB差异"),
            ("模型治理", "100%调用记录版本、哈希、耗时、输出；异常ABSTAIN", "model_runs与模型卡"),
            ("Telegram", "重复发送0；详情深链可达同一event_id", "outbox审计"),
            ("质量", "现有232项不退化；核心覆盖率≥80%；安全审计PASS", "CI/pytest/coverage"),
            ("安全", "密钥泄漏0、交易接口0、Replay伪装0", "secret scan与路由审计"),
        ],
        [1900, 4650, 2810],
        font_size=8.8,
    )

    add_heading(doc, "15. 95+评分映射与诚实判断", 1)
    add_heading(doc, "15.1 校方硬门映射", 2)
    add_table(
        doc,
        ["维度", "V4.0响应", "还必须人工完成"],
        [
            ("高难度认定", "与标准A10区分：证据链、时序身份、最终性、服务器产品、可审计回放", "Day 1当面审批；未通过则理论上最高95"),
            ("过程规范20", "3 Sprint、Git、.agent、AI日志、Review/Retro、五份设计", "每天留真实证据，不补写历史"),
            ("代码理解30", "3个小而纯的手写禁飞区；外围契约化", "逐行解释、Bug修复、即兴修改"),
            ("工程质量20", "VPS、HTTPS、API、Web、账本、恢复、覆盖率、可观测", "运行证据和故障演练"),
            ("功能创新20", "声明—证据—结论、正反例Replay、风险路由与ABSTAIN", "真实指标与消融，不靠宣传"),
            ("团队协作10", "PO/SM/QA/演示职责与贡献矩阵", "2–4人均有代码、评审、测试证据"),
        ],
        [1750, 4580, 3030],
        font_size=8.8,
    )
    add_heading(doc, "15.2 分数判断", 2)
    add_table(
        doc,
        ["状态", "判断", "理由"],
        [
            ("当前今天", "优秀研究底座，但不是95交付态", "有真实数据、证据、审计和232测试；缺Web/VPS/模型/Replay/Git/.agent/coverage"),
            ("V4核心全部实现", "现实区间92–96", "产品、AI、工程和演示闭环，具备高质量自主项目形态"),
            ("再通过高难度审批与强现场考核", "95+可争取", "分数由实际运行、过程证据、个人理解和现场表现共同决定"),
        ],
        [1900, 2500, 4960],
        font_size=9.0,
    )
    add_callout(doc, "最终判断", "这仍然是优秀项目，而且比只用Telegram的版本更像95+作品；但95+不是因为多了一台服务器，而是因为研究底座被转化为可部署、可恢复、可评测、可走查、能现场修改的完整产品。", fill=LIGHT_GREEN, accent=GREEN)

    add_heading(doc, "16. 风险登记与升级门", 1)
    add_table(
        doc,
        ["风险", "预防 / 降级", "停止扩展条件"],
        [
            ("范围过大", "固定四页Web、一个窄模型、两个Replay；其余进Stretch", "任一核心Gate未闭环"),
            ("现场无SEC事件", "LIVE健康 + RECENT真实捕获 + REPLAY主演示", "不得临时伪造LIVE"),
            ("VPS/网络故障", "Windows本地同版本Replay；冻结证据包", "公网不可达立即切本地"),
            ("模型无增益", "Shadow部署、展示诚实指标与ABSTAIN", "不超过规则基线则不参与排序"),
            ("SQLite锁", "单写者、API只读、WAL、短事务", "达到触发条件再迁Postgres"),
            ("禁飞区讲不清", "缩小纯内核、结对讲解、每日手写练习", "讲不清则重写，不扩大功能"),
            ("服务器混入交易程序", "独立Linux用户、目录、容器、网络与secret", "绝不挂载既有量化目录"),
        ],
        [1800, 4800, 2760],
        font_size=8.8,
    )

    add_heading(doc, "17. 最终交付物与Definition of Done", 1)
    add_table(
        doc,
        ["交付物", "必须内容", "状态规则"],
        [
            ("网页终端", "四页、三模式、证据矩阵、时间线、健康与模型页", "公网和本地均可启动"),
            ("服务端", "Caddy、FastAPI、Worker、Streamlit、模型、持久卷、Compose", "重启与健康检查通过"),
            ("数据资产", "正式账本、证据快照、Replay包、备份与恢复记录", "哈希可验、模式隔离"),
            ("AI资产", "Agent Schema、提示词/工具版本、小模型工件、数据卡、模型卡", "指标可复现、异常可降级"),
            ("工程资产", "有效Git、.agent、五份设计、CI、coverage、安全审计", "过程真实、当前可复跑"),
            ("答辩资产", "3分钟脚本、禁飞区讲解、Bug/即兴演练记录、离线Runbook", "每位成员独立完成"),
        ],
        [1750, 4980, 2630],
        font_size=8.8,
    )
    add_heading(doc, "17.1 DoD硬门", 2)
    for item in [
        "任何计划能力都必须有运行截图、日志、测试或数据库记录后才能改写为“已实现”。",
        "Web、API、Worker、账本、Replay、模型和Telegram共享同一event_id与trace_id。",
        "NINEQ正例与WOLF2反例均能从冻结输入复现，且LIVE/REPLAY不可混淆。",
        "核心覆盖率≥80%，232项既有测试不退化，安全审计和secret scan通过。",
        "VPS重启、容器重启、来源超时、模型加载失败、Telegram重试均有可演示降级。",
        "3个禁飞区、3轮Bug注入和每人即兴修改均有真实记录且能现场解释。",
    ]:
        add_list(doc, item, bullet_id)

    add_heading(doc, "18. 立项结论", 1)
    add_para(doc, "V4.0把Finance Radar从“Telegram上的财经提醒”重新定义为“服务器部署的证据型金融事件情报工作台”。它保留现有最强资产——真实事件历史、修订、证据、时序身份和拒绝案例——同时补上主展示终端、持续运行、历史恢复、小模型治理和确定性回放。该路线与普通新闻聚合Agent有清晰差异，也能在12天内形成完整、可验收的垂直切片。")
    add_para(doc, "项目最值得强调的不是“模型预测更准”，而是“系统更会证明、更会拒绝、更能保存历史，也更能在现场被解释和修改”。只要严格执行核心Gate，不被可选技术分散精力，这个版本足以进入优秀档；再叠加高难度审批、真实过程证据和扎实现场表现，95+具有现实可行性。")
    add_callout(doc, "最终取舍", "先完成可部署、可恢复、可回放的一条完整事件链；再增加来源和模型。先让评委看懂证据与拒绝；再展示速度和界面。", fill=LIGHT_BLUE)

    add_heading(doc, "19. 主要依据与参考资料", 1)
    add_reference(doc, 1, "北京林业大学2026实训动员会（本地课程文件）")
    add_reference(doc, 2, "北京林业大学理学院2026年实现项目列表（本地课程文件）")
    add_reference(doc, 3, "FastAPI Deployment Concepts", "https://fastapi.tiangolo.com/deployment/concepts/")
    add_reference(doc, 4, "FastAPI WebSockets", "https://fastapi.tiangolo.com/advanced/websockets/")
    add_reference(doc, 5, "Docker Compose in Production", "https://docs.docker.com/compose/how-tos/production/")
    add_reference(doc, 6, "SQLite Online Backup API", "https://www.sqlite.org/backup.html")
    add_reference(doc, 7, "SQLite Write-Ahead Logging", "https://www.sqlite.org/wal.html")
    add_reference(doc, 8, "SEC Developer Resources and Fair Access", "https://www.sec.gov/about/developer-resources")
    add_reference(doc, 9, "NIST AI 600-1 Generative AI Profile", "https://nvlpubs.nist.gov/nistpubs/ai/NIST.AI.600-1.pdf")
    add_reference(doc, 10, "W3C PROV-O: The PROV Ontology", "https://www.w3.org/TR/prov-o/")
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
