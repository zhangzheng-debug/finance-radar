from __future__ import annotations

import json
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    docx = root / "financial_event_radar_project_proposal_v2_1_human.docx"
    markdown = root / "financial_event_radar_project_plan_v2_0_ai.md"

    report: dict[str, object] = {}
    with ZipFile(docx) as archive:
        bad_part = archive.testzip()
        names = set(archive.namelist())
        report["zip_integrity"] = "PASS" if bad_part is None else f"FAIL:{bad_part}"
        report["required_parts"] = all(
            part in names
            for part in (
                "[Content_Types].xml",
                "word/document.xml",
                "word/styles.xml",
                "word/numbering.xml",
            )
        )

    document = Document(docx)
    chunks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    human_text = "\n".join(chunks)
    ai_text = markdown.read_text(encoding="utf-8")

    report["human"] = {
        "bytes": docx.stat().st_size,
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "sections": len(document.sections),
        "headings": sum(
            1
            for paragraph in document.paragraphs
            if paragraph.style and paragraph.style.name.startswith("Heading")
        ),
    }
    report["ai"] = {
        "bytes": markdown.stat().st_size,
        "lines": ai_text.count("\n") + 1,
        "balanced_fences": ai_text.count("```") % 2 == 0,
    }

    report["shared_markers"] = {
        token: {
            "human": token.lower() in human_text.lower(),
            "ai": token.lower() in ai_text.lower(),
        }
        for token in (
            "M1",
            "OpenNews",
            "NewsLiquid",
            "SQLite WAL",
            "provider_assessment",
            "Telegram",
            "IBKR",
            "Twelve Data",
            "SEC",
            "不交易",
            "Schema",
            "69",
            "562",
            "781",
        )
    }
    report["placeholders"] = {
        token: {
            "human": token.lower() in human_text.lower(),
            "ai": token.lower() in ai_text.lower(),
        }
        for token in ("TODO", "TBD", "PLACEHOLDER", "待补充", "Lorem ipsum")
    }
    report["empty_tables"] = sum(
        1
        for table in document.tables
        if not any(cell.text.strip() for row in table.rows for cell in row.cells)
    )
    report["empty_heading_count"] = sum(
        1
        for paragraph in document.paragraphs
        if paragraph.style
        and paragraph.style.name.startswith("Heading")
        and not paragraph.text.strip()
    )
    report["page_size_letter"] = all(
        round(section.page_width.inches, 2) == 8.5
        and round(section.page_height.inches, 2) == 11.0
        for section in document.sections
    )

    expected_width = 9360
    grid_violations = []
    row_violations = []
    for table_index, table in enumerate(document.tables):
        grid_widths = [int(column.get(qn("w:w"))) for column in table._tbl.tblGrid]
        if sum(grid_widths) != expected_width:
            grid_violations.append({"table": table_index, "sum": sum(grid_widths)})
        for row_index, row in enumerate(table._tbl.tr_lst):
            widths = []
            for cell in row.tc_lst:
                width_node = cell.tcPr.find(qn("w:tcW"))
                if width_node is not None and width_node.get(qn("w:w")):
                    widths.append(int(width_node.get(qn("w:w"))))
            if widths and sum(widths) != expected_width:
                row_violations.append(
                    {"table": table_index, "row": row_index, "sum": sum(widths)}
                )
    report["table_geometry"] = {
        "expected_width_dxa": expected_width,
        "grid_violations": grid_violations,
        "row_violations": row_violations,
        "pass": not grid_violations and not row_violations,
    }

    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
