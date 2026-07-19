#!/usr/bin/env python3
"""Build the verified-baseline Finance Radar V5.2 human proposal."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_finance_radar_plan_v3 import (
    BLUE,
    DARK_BLUE,
    GOLD,
    GREEN,
    LIGHT,
    LIGHT_BLUE,
    LIGHT_GOLD,
    LIGHT_GREEN,
    LIGHT_RED,
    MUTED,
    NAVY,
    RED,
    add_callout,
    add_heading,
    add_list,
    add_num_defs,
    add_page_number,
    add_para,
    add_reference,
    add_table,
    configure_styles,
    set_paragraph_spacing,
    set_run_font,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "financial_event_radar_project_proposal_v5_2_human.docx"


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    configure_styles(doc)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    set_paragraph_spacing(header, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_run_font(
        header.add_run("FINANCE RADAR  ·  EVIDENCE TERMINAL BASELINE  ·  V5.2"),
        size=8.5,
        bold=True,
        color=MUTED,
    )
    footer = section.footer.paragraphs[0]
    set_paragraph_spacing(footer, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_run_font(
        footer.add_run("北京林业大学理学院2026实训  ·  研究监控 / 无交易能力  ·  "),
        size=8.5,
        color=MUTED,
    )
    add_page_number(footer)


def add_display_para(
    doc: Document,
    text: str,
    *,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
    before: float = 0,
    after: float = 8,
    line: float = 1.0,
    size: float = 11,
    color=NAVY,
    bold: bool = False,
    italic: bool = False,
):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=before, after=after, line=line, align=align)
    set_run_font(
        paragraph.add_run(text),
        size=size,
        color=color,
        bold=bold,
        italic=italic,
    )
    return paragraph


def add_cover(doc: Document) -> None:
    add_display_para(
        doc,
        "北京林业大学理学院 2026 实训 · 自主高难度创新项目",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=34,
        after=18,
        line=1.0,
        size=10.5,
        bold=True,
        color=BLUE,
    )
    add_display_para(
        doc,
        "基于多源证据链与实时 Web 控制台的",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=5,
        line=1.05,
        size=20,
        bold=True,
        color=NAVY,
    )
    add_display_para(
        doc,
        "金融事件情报 Agent",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=10,
        line=1.05,
        size=30,
        bold=True,
        color=NAVY,
    )
    add_display_para(
        doc,
        "全极性事实核验 · 下行风险优先路由 · 历史账本 · 确定性回放",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=22,
        line=1.1,
        size=12.5,
        color=DARK_BLUE,
    )
    add_table(
        doc,
        ["版本", "基线日期", "主展示", "部署形态"],
        [["V5.2 持续取证任务书", "2026-07-19", "Evidence Terminal", "新加坡 VPS / HTTPS"]],
        [2100, 1800, 2760, 2700],
        font_size=9.2,
        center_cols=(0, 1, 2, 3),
    )
    add_callout(
        doc,
        "一句话价值",
        "把外部消息变成可复核、可拒绝、可回放的金融事件证据链；系统不预测涨跌，不读取资金、仓位或交易账户，也不执行交易。",
        fill=LIGHT_BLUE,
    )
    add_callout(
        doc,
        "版本说明",
        "V5.2 已完成22路分级来源、证据终端、官方HTML/PDF/JSON不可变存档、每5分钟自动采集、每轮最多4份原始证据、事件证据无TTL、30日/12周在线备份、回放、模型盲测、加密异机全量恢复、24小时哈希链取证机制和课程真实性机器门禁。完整24小时窗口仍由机器保持等待；外部盲测诚实失败并阻断模型晋级；教师/学生/Git/练习证据缺失时审计器保持NOT_READY，不允许人工包装。",
        fill=LIGHT_GOLD,
        accent=GOLD,
    )
    add_display_para(
        doc,
        "自主高难度创新选题申报稿 / 95+ 冲刺版",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=14,
        after=0,
        line=1.0,
        size=10,
        italic=True,
        color=MUTED,
    )
    doc.add_page_break()


def build() -> Path:
    doc = Document()
    configure_document(doc)
    bullet_id, decimal_id = add_num_defs(doc)
    props = doc.core_properties
    props.title = "基于多源证据链与实时Web证据终端的金融事件情报Agent V5.2"
    props.subject = "北京林业大学理学院2026实训自主高难度创新项目任务书"
    props.author = "Finance Radar Project Team"
    props.keywords = "Finance Radar, Evidence Agent, Web, Replay, Audit, No Trading"
    add_cover(doc)

    add_heading(doc, "执行摘要", 1)
    add_para(
        doc,
        "Finance Radar 已从新闻聚合设想升级为公网运行的金融事件证据终端。它持续接入官方与发现型来源，保存原始观测、修订、事件版本和精确证据段落；通过三栏 Web 工作面同时呈现事件流、声明—证据矩阵和判断上下文；通过固定回放消除“答辩现场恰好没有 SEC 重大事件”的随机性。Telegram 只承担通知和网页深链接，不再是唯一终端。",
    )
    add_callout(
        doc,
        "当前判断",
        "工程与产品成品已经具备优秀档基础，92–96 分是现实竞争区间。95+ 可以争取，但剩余决定性因素已经不是再堆页面，而是教师高难度认定、学生手写禁飞区、真实分工与提交历史、24 小时运行证据，以及现场 Bug 修复和即兴修改表现。",
        fill=LIGHT_GREEN,
        accent=GREEN,
    )

    add_heading(doc, "1. 已验证基线与诚实缺口", 1)
    add_heading(doc, "1.1 公网成品快照", 2)
    add_table(
        doc,
        ["板块", "2026-07-19 实测证据", "状态"],
        [
            ("公网入口", "HTTPS Web 与 FastAPI；TLS 1.3；API/Web 仅在 loopback 监听", "通过"),
            ("数据账本", "22来源；1194事件；3951原始观测；2117版本；2394证据；Schema 12", "通过"),
            ("运行服务", "release 20260719T044852Z；API/Web/Worker/Backup/本地模型五服务 active；Worker每5分钟自动采集", "通过·24h等待"),
            ("Evidence Agent", "claims/edges/summary；81份官方原始HTML/PDF快照与2份精确引文；每轮最多新增4份；分页绕过已存档/失败头部；SHA-256失败0", "通过·持续积累"),
            ("小模型", "V1外部盲测失败；V2候选拒绝；生产保持Shadow；V3盲标24条任务仍为0人审", "受治理·Shadow"),
            ("备份恢复", "9860项清单全过；Schema 12/3完整恢复；1559.8MB全服务准备；在线保留30日/12周；第二恢复密钥ACL独占", "通过"),
            ("质量", "360 tests + 17 subtests；公网产品19/19、行情17/17；五页AppTest通过；上一版浏览器6/6与可访问性纯PASS，本次UI待刷新", "通过·视觉待刷新"),
            ("故障与负载", "6项故障注入全过；120请求/15并发成功率100%；公网p95 0.832秒；来源故障后自动恢复", "通过"),
        ],
        [1800, 5900, 1660],
        font_size=8.6,
        center_cols=(2,),
        status_col=2,
    )
    add_heading(doc, "1.2 仍未完成，禁止包装", 2)
    for item in [
        "V3双人盲标基础设施已完成，但24条真实任务仍为0人审；不得由AI或同一人伪造双审标签。",
        "连续24小时运行窗口尚在积累；本地计划任务每15分钟自动保存哈希链接证据；只在服务端完整门槛自然满足后转绿，不能提前包装为完成。",
        "三个禁飞区必须由学生本人手写、测试、提交和讲解，现有 AI 代码不能倒签为手写。",
        "教师对高难度自选题和第三禁飞区的批准属于外部课程门槛。",
        "团队分工、个人提交、盲测和即兴修改必须留下真实过程证据，不能补写历史。",
        "当前行情UI改版已通过结构与公网API验收，但大屏、桌面、移动视觉/交互/可访问性矩阵仍须对新release刷新，旧截图只作为基线。",
    ]:
        add_list(doc, item, bullet_id)

    add_heading(doc, "2. 项目定位、用户与边界", 1)
    add_para(
        doc,
        "目标用户是需要快速理解重大公司、监管和市场结构事件的研究者。系统回答四类问题：发生了什么；哪些原子声明有一手证据；哪些材料支持、反驳或仍不足；该事件是否需要优先人工审核。它不是资讯门户、聊天机器人、量化交易系统或收益预测器。",
    )
    add_table(
        doc,
        ["语义层", "回答的问题", "允许输出"],
        [
            ("事实极性", "事件内容是正面、负面、中性还是混合？", "FAVORABLE / ADVERSE / NEUTRAL / MIXED"),
            ("证据状态", "声明是否有精确引用，是否冲突？", "EVIDENCE_READY / INSUFFICIENT / HUMAN_REVIEW"),
            ("风险路由", "是否值得优先人工复核？", "RISK_REVIEW / NON_TARGET / ABSTAIN"),
            ("事件最终性", "法律或业务阶段是否已达到最终状态？", "由确定性状态机与人工规则控制"),
        ],
        [1900, 4100, 3360],
        font_size=9,
    )
    add_callout(
        doc,
        "为何强调负面事件",
        "重大下行风险往往具有损失不对称和审核时效性，因此模型优先路由负面高风险事件是合理的产品目标；这不等于只采集负面新闻，更不等于负面事件一定导致价格下跌。正面和中性事件继续进入账本，并作为 NON_TARGET 对照。",
        fill=LIGHT_GOLD,
        accent=GOLD,
    )

    add_heading(doc, "3. 成品架构与数据流", 1)
    add_table(
        doc,
        ["层", "组件", "已落地职责"],
        [
            ("边缘", "Nginx + Let's Encrypt", "公网 8443 HTTPS、反向代理、证书自动续期"),
            ("展示", "Streamlit Evidence Terminal", "紧凑状态壳层、三栏事件工作台、回放时间轴、运维 SLO"),
            ("契约", "FastAPI", "只读查询、受控 Replay/Agent/人工覆盖；统一 envelope 与 trace_id"),
            ("处理", "systemd Worker", "持续采集、证据丰富、候选整理、shadow 推理、outbox"),
            ("数据", "SQLite WAL + SHA-256对象仓", "Schema 12正式账本；Schema 3运维/Replay/备份/盲标；官方HTML/PDF原始快照"),
            ("通知", "Telegram outbox", "幂等、租约、默认 dry-run、显式 --send、网页深链"),
            ("恢复", "在线备份 + 加密异机同步", "SSH/SHA/AES-GCM；防路径穿越；全清单核对；两库隔离恢复与明文清理"),
        ],
        [1300, 2600, 5460],
        font_size=8.8,
    )
    add_heading(doc, "3.1 来源分层：事件源与行情源不混写", 2)
    add_table(
        doc,
        ["类别", "当前来源", "用途与边界"],
        [
            ("P0 监管/公共安全", "SEC、CFTC、FTC、FDIC、FDA", "一手规则、执法、申报与警示；可核验事件事实"),
            ("P0 宏观/政策", "Federal Reserve、BLS、ECB、EIA", "政策、通胀、劳工与能源发布；不自动推导资产方向"),
            ("P1 发行人官方", "NVIDIA newsroom 等注册官方渠道", "公司声明与修订；权威级别不等于正面或负面"),
            ("P2 发现", "OpenNews、历史研究候选", "扩大召回；只能触发核验，不能自动形成最终事实"),
            ("行情观察", "Binance公共现货 + Twelve Data落库；IBKR TWS仅本机探针", "首次真实观察为基线；T+5m/30m/1d到期采集；错过则MISSED且不补最新价；收益仅作事件后审计"),
        ],
        [1900, 3300, 4160],
        font_size=8.6,
    )
    add_heading(doc, "3.2 安全隔离", 2)
    for item in [
        "Finance Radar 使用独立目录、Linux 用户、端口、systemd 服务和数据库；",
        "不进入、不修改、不重启 `/root/ethusdc-pivot-bot`，不读取其中任何交易凭据；",
        "API 没有 orders、positions、balances、brokerage_accounts 或 trade_execution 路由；",
        "行情只读且只用于时间对齐和事后审计，严禁作为事实真假或因果收益的证明。",
    ]:
        add_list(doc, item, bullet_id)

    add_heading(doc, "4. Web Situation Room 与确定性演示", 1)
    add_table(
        doc,
        ["页面", "核心信息", "答辩证明"],
        [
            ("Situation Room", "紧凑状态带、快捷流、全终端检索、数据驱动命令条、来源健康、Worker、备份和审核队列", "系统在线且可快速定位事件"),
            ("Event Intelligence", "只读Facets、事件族联想、来源筛选、URL连续性、本机命名Flow、声明—证据矩阵", "同屏发现、判断、核验"),
            ("Replay Lab", "四固定案例、横向时间轴、证据变化、官方更正、决策与告警资格", "不依赖随机外部事件"),
            ("Operations & Model", "SLO、异常来源、备份、原始证据存档、模型卡、漂移门与安全计数", "工程运行与治理可审计"),
            ("Adjudication Studio", "双审核者互盲、隐藏模型/来源/事后行情、冲突第三人裁决", "证明训练标签来自人审内容而非来源捷径"),
        ],
        [2100, 4400, 2860],
        font_size=8.9,
    )
    add_heading(doc, "4.1 三模式协议", 2)
    for item in [
        "LIVE：证明外部连接、来源游标和最新成功采集；没有新事件也可判定成功。",
        "RECENT_CAPTURE：展示真实历史中的发布时间、接收时间、处理时间、原文和哈希。",
        "REPLAY：用冻结真实案例证明证据门、冲突处理、最终性与模型路由可重复。",
    ]:
        add_list(doc, item, decimal_id)
    add_callout(
        doc,
        "三分钟主演示",
        "Situation Room 健康 → 最近真实事件 → SEC 破产案例从 P2 发现到 P0 一手证据 → 冲突/正面控制 → 模型与数据卡 → 备份、trace 和无交易证明。",
        fill=LIGHT_GREEN,
        accent=GREEN,
    )

    add_heading(doc, "5. Evidence Agent：会证明，也会拒绝", 1)
    add_para(
        doc,
        "Agent 采用四节点结构：声明抽取、证据计划、证据关系提议、带引文摘要。所有输出强制结构化，证据边必须包含精确段落、来源 URL、权威层级和内容对象 SHA-256；审核过的官方证据还保存原始HTML/PDF/JSON字节。快照器只接收注册官方来源；其HTTP链接先安全升级HTTPS，重定向后再次复核域名；单件限制10MiB、每轮最多新增4份，事件与证据不设TTL。缺证据返回 INSUFFICIENT，未解决冲突返回 HUMAN_REVIEW，原始快照也不能自动核验事实或进入模型特征。",
    )
    add_table(
        doc,
        ["对象", "关键字段", "审计价值"],
        [
            ("EventClaim", "claim_id、text、event_version、material、verification_state", "把长新闻拆成可逐条核验的声明"),
            ("EvidenceEdge", "relation、exact_excerpt、source_url、tier、object_sha256", "证明哪一段支持或反驳哪条声明"),
            ("AgentDecision", "trace_id、provider、snapshot、tool_calls、guardrails、latency", "完整回放模型/工具/守卫状态"),
            ("HumanOverride", "actor、reason、before、after、time", "人工决定可追溯，不静默覆盖"),
        ],
        [1800, 4700, 2860],
        font_size=8.8,
    )
    add_callout(
        doc,
        "当前诚实口径",
        "结构化 Agent 与回环 Qwen2.5-0.5B/llama.cpp 已在线，真实SEC事件返回 llm_used=true；小模型只生成有引用的advisory summary。声明关系、证据门和最终状态仍由确定性图控制，非法输出自动回退，晋级结论保持REMAIN_SHADOW。",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "6. 小模型：重大下行风险人工审核路由", 1)
    add_table(
        doc,
        ["项目", "V5.2 实际契约"],
        [
            ("任务", "重大下行风险审核路由，不是涨跌或收益预测"),
            ("模型", "word/char TF-IDF + 校准、类别平衡逻辑回归；CPU shadow"),
            ("样本", "897；RISK_REVIEW 546，NON_TARGET 351；包含正面/中性/拒绝控制"),
            ("切分", "时间优先的发行人/事件链连通分组；两类重叠计数均为0"),
             ("分组留出", "覆盖率82.7%；弃权17.3%；覆盖样本准确率95.7%"),
             ("外部盲测", "40条零重叠；风险召回100%；正常新闻误报95%；预注册门槛FAIL"),
             ("治理", "模型/数据卡、SHA-256、三组消融、漂移门、捷径系数审计；强制REMAIN_SHADOW"),
            ("禁止", "LONG、SHORT、目标价、预期收益、严重度、自动告警许可"),
        ],
        [2000, 7360],
        font_size=9,
    )
    add_para(
        doc,
        "模型仍保持 Shadow Mode。分组留出集上的三组消融表现较好，但 label-first 外部盲测揭示了真实跨来源失效：39/40条被路由为风险，正常官方新闻误报率95%。系数审计进一步发现 event_family、event_type、discovery_source 等内部分类语言形成捷径，且 NON_TARGET 训练样本不是普通新闻的代表。该失败不回填标签、不在 blind-v1 上调参；V1 只做候选队列高召回路由，不能包装成正负新闻总分类器。",
    )

    add_heading(doc, "7. 木桶式剩余实施计划", 1)
    add_table(
        doc,
        ["优先级", "短板", "完成定义"],
        [
            ("P0", "教师审批与禁飞区", "真实证据路径/哈希/提交写入manifest；3个学生内核完成后审计转绿"),
            ("P0", "真实过程与现场能力", "成员角色、每人3次Bug和1次即兴修改；当前11项机器检查保持false"),
            ("P1", "24小时运行证据", "每15分钟SHA-256链已自动取证；等待服务端完整门槛后生成PASS报告"),
            ("DONE", "本地证据摘要模型", "Qwen2.5 0.5B / llama.cpp回环部署；冻结对比与线上接受已过；仅做证据摘要，不决定标签"),
            ("DONE", "异机备份与换机", "AES-GCM；9860项清单；1559.8MB全服务准备；在线保留30日/12周；新VPS失败前置门；显式激活；仓库外第二密钥"),
            ("DONE", "原始证据存档", "81份官方HTML/PDF快照；注册官方来源安全升级HTTPS；每轮最多4份；分页绕过已存档/失败头部；支持HTML/PDF/JSON；无TTL；哈希抽检失败0"),
            ("DONE", "模型诚实补证", "40条外部盲测已冻结并执行；失败归因与晋级阻断已落地；继续 Shadow"),
            ("OPTION", "模型V2", "只在时间允许时做内容字段去泄漏和独立开发负样本；锁定后必须新建blind-v2"),
            ("DONE", "安全与体验", "限流/负载已过；五页结构、键盘导航和可执行断网终端已验收；当前UI浏览器矩阵仍待刷新"),
        ],
        [1200, 2500, 5660],
        font_size=8.7,
    )
    add_callout(
        doc,
        "停止扩展规则",
        "只要 P0/P1 有红项，就不投入 PostgreSQL、React、Kubernetes、Redis、Celery 或任何交易功能。每完成一个短板，都必须留下可复核证据，再进入下一桶板。",
        fill=LIGHT_GOLD,
        accent=GOLD,
    )

    add_heading(doc, "8. 学生禁飞区与现场考核", 1)
    add_table(
        doc,
        ["禁飞区", "建议职责", "学生现场证明"],
        [
            ("event_fingerprint.py", "事件归一、修订与同一性判断", "标准化、哈希、误合并/漏合并边界；≥18项测试"),
            ("evidence_gate.py", "声明覆盖、冲突和来源独立性", "逐声明聚合、fail closed；≥18项测试"),
            ("finality_gate.py", "阶段状态、身份冲突、最终性与等级上限", "状态机、不变量、非法转换；≥20项测试并获教师批准"),
        ],
        [2300, 3600, 3460],
        font_size=8.8,
    )
    add_para(
        doc,
        "禁飞区必须无网络、无数据库、无 LLM，控制在约80–150有效行。学生保留设计草图、失败测试、手写提交和讲解脚本。AI 只允许协助外围适配、文档和测试基础设施，不能代写最终禁飞区。",
    )
    add_heading(doc, "8.1 Bug 注入与即兴修改", 2)
    for item in [
        "每个 Sprint 由非实现者盲注3个 Bug；内部题库必须标为练习，不冒充官方题。",
        "每位成员至少完成3次计时练习，30分钟优秀目标为修复并解释2–3个问题。",
        "即兴修改优先选择字段、筛选器、阈值、Replay fixture 或来源适配器，必须补回归测试。",
        "答题话术固定为：复现 → 定位证据 → 根因 → 最小补丁 → 回归 → 安全边界。",
    ]:
        add_list(doc, item, bullet_id)

    add_heading(doc, "9. 定量验收与交付物", 1)
    add_table(
        doc,
        ["验收项", "通过标准", "当前 / 剩余"],
        [
            ("公网产品", "HTTPS Web/API 可用；五页和三模式可操作", "已通过"),
            ("证据治理", "重大声明有精确引文；缺失/冲突强制拒绝或人工复核", "已通过安全后备"),
            ("运行恢复", "服务自启；24小时无重复爆炸；加密包清单/两库/发布版隔离恢复成功", "全量恢复已过；24小时待积累"),
            ("测试", "全量<60秒；禁飞区行/分支覆盖>80%；定向测试<10秒", "全量已过；禁飞区待学生完成"),
            ("安全", "secret 泄漏0、交易路由0、模型越权0、Replay伪装0", "已通过"),
            ("过程", "真实 Git、分工、Review/Retro、AI 使用记录", "真实性审计已建；当前0提交/11项缺证，待学生形成"),
            ("现场", "每人讲解负责内核；Bug 与即兴修改在时限内完成", "待团队演练"),
        ],
        [1800, 4560, 3000],
        font_size=8.7,
    )
    add_heading(doc, "9.1 最终交付包", 2)
    # Start a fresh numbering sequence for this independent deliverables list.
    # Reusing ``decimal_id`` would make Word continue the earlier three-mode
    # protocol and render this section as items 4-11.
    _, deliverable_decimal_id = add_num_defs(doc)
    for item in [
        "公网 Web/API、systemd 服务、Nginx HTTPS 和回滚 release；",
        "Schema 12 正式账本、Schema 3 运维/Replay/备份/盲标轨迹、官方HTML/PDF与精确引文内容寻址对象；",
        "四案例 Replay（含SEC官方更正撤回）、Agent trace、模型/数据卡与结构化训练清单；",
        "30日/12周备份、每日 AES-GCM 异机副本、全清单/两库隔离恢复报告、24小时哈希链、故障注入和120/15负载报告；",
        "两层离线交付：110份精选证据ZIP用于无服务复核；可执行终端含22条真实事件、证据、Replay、影子模型、API和五页Web；只允许回环，不含采集器、Telegram、券商/交易所客户端、密钥或交易能力；",
        "人读任务书、AI 执行规范、.agent 过程文件和答辩 Runbook；",
        "可直接填写的教师审批单与学生执行证据包（不冒充真实审批/练习）；",
        "学生手写禁飞区、个人贡献、计时练习和教师审批材料，以及READY课程审计报告。",
    ]:
        add_list(doc, item, deliverable_decimal_id)

    add_heading(doc, "10. 95+判断与立项结论", 1)
    add_table(
        doc,
        ["层次", "判断", "决定因素"],
        [
            ("工程成品", "已经优秀", "真实公网闭环、证据历史、回放、模型治理、备份恢复和故障证据"),
            ("课程项目", "具备92–96竞争力", "任务难度、完整度和答辩确定性已足够"),
            ("最终95+", "可以争取但不保证", "教师认定、学生独立理解、真实过程与现场修复/修改"),
        ],
        [1800, 2600, 4960],
        font_size=9,
    )
    add_callout(
        doc,
        "最终取舍",
        "后续最有价值的工作不是继续堆功能，而是把当前成品变成学生真正掌握、教师认可、连续运行、现场能改的作品。完成 P0/P1 后即可封版；其余增强全部降为可选。",
        fill=LIGHT_GREEN,
        accent=GREEN,
    )
    add_para(
        doc,
        "Finance Radar 的突出点不在于“模型告诉你买卖”，而在于系统能够保存消息如何变化、指出哪条声明由哪段一手证据支持、在冲突或不足时拒绝下结论，并把数据、模型、Agent、备份和演示统一在可审计的工程链路中。这一定位与普通新闻聚合或情绪分析项目有明确差异，适合作为自主高难度创新选题。",
    )

    add_heading(doc, "主要依据与参考资料", 1)
    add_reference(doc, 1, "北京林业大学2026实训动员会（课程文件）")
    add_reference(doc, 2, "北京林业大学理学院2026年实现项目列表（课程文件）")
    add_reference(doc, 3, "SEC Developer Resources and Fair Access", "https://www.sec.gov/about/developer-resources")
    add_reference(doc, 4, "ECB RSS Feeds", "https://www.ecb.europa.eu/home/html/rss.en.html")
    add_reference(doc, 5, "U.S. EIA RSS Feeds", "https://www.eia.gov/tools/rssfeeds/")
    add_reference(doc, 6, "NVIDIA Newsroom RSS", "https://nvidianews.nvidia.com/releases.xml")
    add_reference(doc, 7, "FastAPI Deployment Concepts", "https://fastapi.tiangolo.com/deployment/concepts/")
    add_reference(doc, 8, "SQLite Online Backup API", "https://www.sqlite.org/backup.html")
    add_reference(doc, 9, "Model Cards for Model Reporting", "https://arxiv.org/abs/1810.03993")
    add_reference(doc, 10, "Datasheets for Datasets", "https://arxiv.org/abs/1803.09010")
    add_reference(doc, 11, "W3C PROV-O", "https://www.w3.org/TR/prov-o/")

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
