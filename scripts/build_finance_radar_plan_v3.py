#!/usr/bin/env python3
"""Build the human-readable Finance Radar V3.0 high-difficulty proposal."""

from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "financial_event_radar_project_proposal_v3_0_human.docx"
ASSET_DIR = ROOT / "tmp" / "plan_v3_assets"
ARCHITECTURE_IMAGE = ASSET_DIR / "finance_radar_v3_architecture.png"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1E293B"
MUTED = "667085"
LIGHT = "F4F6F9"
LIGHT_BLUE = "EAF2F8"
LIGHT_GOLD = "FFF7E6"
LIGHT_GREEN = "EDF7F0"
LIGHT_RED = "FDECEC"
WHITE = "FFFFFF"
GRID = "D7DEE8"
GREEN = "217346"
GOLD = "9A6700"
RED = "9B1C1C"


def set_run_font(run, *, size=11, bold=None, italic=None, color=INK, name="Calibri"):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name)
    rpr.rFonts.set(qn("w:hAnsi"), name)
    rpr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, *, before=0, after=8, line=1.333, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.widow_control = True
    if align is not None:
        paragraph.alignment = align


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=GRID, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths_dxa):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {CONTENT_WIDTH_DXA}: {widths_dxa}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        existing = tbl_pr.find(qn(tag))
        if existing is not None:
            tbl_pr.remove(existing)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            set_cell_border(cell)


def add_num_defs(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abs = max(abstract_ids, default=0) + 1
    next_num = max(num_ids, default=0) + 1

    def create(kind, abs_id, num_id):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abs_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "280")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "290")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Calibri")
        fonts.set(qn("w:hAnsi"), "Calibri")
        fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r_pr.append(fonts)
        lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abs_id))
        num.append(abstract_ref)
        numbering.append(num)

    create("bullet", next_abs, next_num)
    create("decimal", next_abs + 1, next_num + 1)
    return next_num, next_num + 1


def set_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def add_para(doc, text="", *, bold_prefix=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             before=0, after=8, line=1.333, size=11, color=INK, keep=False):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=before, after=after, line=line, align=align)
    p.paragraph_format.keep_together = keep
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, size=size, bold=True, color=color)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest, size=size, color=color)
    else:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color)
    return p


def add_list(doc, text, num_id, *, size=11, color=INK):
    p = doc.add_paragraph()
    set_num(p, num_id)
    set_paragraph_spacing(p, before=0, after=4, line=1.208, align=WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    p.add_run(text)
    return p


def add_table(doc, headers, rows, widths_dxa, *, header_fill=LIGHT, font_size=9.5,
              center_cols=(), status_col=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, header in enumerate(headers):
        cell = hdr.cells[idx]
        shade_cell(cell, header_fill)
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, before=0, after=0, line=1.05, align=WD_ALIGN_PARAGRAPH.CENTER)
        run = p.add_run(header)
        set_run_font(run, size=9.5, bold=True, color=NAVY)
    for values in rows:
        row = table.add_row()
        for idx, value in enumerate(values):
            p = row.cells[idx].paragraphs[0]
            align = WD_ALIGN_PARAGRAPH.CENTER if idx in center_cols else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(p, before=0, after=0, line=1.08, align=align)
            text = str(value)
            color = INK
            bold = False
            if status_col == idx:
                upper = text.upper()
                if "已实现" in text or upper == "PASS":
                    color, bold = GREEN, True
                    shade_cell(row.cells[idx], LIGHT_GREEN)
                elif "计划" in text or "待" in text:
                    color, bold = GOLD, True
                    shade_cell(row.cells[idx], LIGHT_GOLD)
                elif "风险" in text or upper == "FAIL":
                    color, bold = RED, True
                    shade_cell(row.cells[idx], LIGHT_RED)
            run = p.add_run(text)
            set_run_font(run, size=font_size, color=color, bold=bold)
    set_table_geometry(table, widths_dxa)
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=2, line=1.0)
    return table


def add_callout(doc, label, text, *, fill=LIGHT_BLUE, accent=BLUE):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=1, after=1, line=1.18, align=WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "22")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), accent)
    borders.append(left)
    p_pr.append(borders)
    r1 = p.add_run(f"{label}  ")
    set_run_font(r1, size=10.5, bold=True, color=accent)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=INK)
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, after=4, line=1.0)
    return p


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, size=9, color=MUTED)


def add_hyperlink(paragraph, text, url, *, color=BLUE, underline=True):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.append(r_fonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    set_paragraph_spacing(hp, before=0, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.LEFT)
    run = hp.add_run("FINANCE RADAR  ·  HIGH-DIFFICULTY PROJECT PROPOSAL V3.0")
    set_run_font(run, size=8.5, bold=True, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    set_paragraph_spacing(fp, before=0, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    r = fp.add_run("北京林业大学理学院2026实训  ·  研究与监控用途  ·  ")
    set_run_font(r, size=9, color=MUTED)
    add_page_number(fp)


def rounded_rectangle(draw, box, radius, fill, outline, width=3):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def make_architecture_diagram(path):
    path.parent.mkdir(parents=True, exist_ok=True)
    width, height = 1800, 1040
    image = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    font_path = r"C:\Windows\Fonts\msyh.ttc"
    bold_path = r"C:\Windows\Fonts\msyhbd.ttc"
    title_font = ImageFont.truetype(bold_path, 48)
    box_title = ImageFont.truetype(bold_path, 29)
    body_font = ImageFont.truetype(font_path, 23)
    small_font = ImageFont.truetype(font_path, 20)

    draw.text((60, 42), "财经事件雷达 V3.0：证据驱动 Agent 总体架构", font=title_font, fill="#17365D")
    draw.text((62, 105), "模型负责语义工作；确定性代码负责身份、最终性、安全门与不可交易边界", font=body_font, fill="#667085")

    boxes = [
        ((70, 205, 320, 520), "多源发现", ["P0/P1 官方源", "P2 聚合新闻", "P3 社交线索", "只读行情源"], "#EAF2F8", "#2E74B5"),
        ((380, 205, 660, 520), "不可变数据层", ["原始快照", "内容哈希", "来源独立组", "事件时间戳"], "#F4F6F9", "#667085"),
        ((720, 170, 1120, 555), "Evidence Agent", ["原子声明抽取", "证据检索规划", "权威工具调用", "支持/反驳/不足", "带引用摘要"], "#FFF7E6", "#9A6700"),
        ((1180, 205, 1480, 520), "确定性安全门", ["时点证券身份", "法律阶段规则", "来源等级约束", "禁止自动升S", "无交易工具"], "#FDECEC", "#9B1C1C"),
        ((1540, 205, 1740, 520), "输出", ["事件卡片", "证据时间线", "Dashboard", "Telegram"], "#EDF7F0", "#217346"),
    ]
    for box, title, lines, fill, outline in boxes:
        rounded_rectangle(draw, box, 24, fill, outline, 4)
        x1, y1, x2, y2 = box
        tw = draw.textbbox((0, 0), title, font=box_title)[2]
        draw.text(((x1 + x2 - tw) / 2, y1 + 25), title, font=box_title, fill=outline)
        yy = y1 + 92
        for line in lines:
            draw.text((x1 + 24, yy), "• " + line, font=small_font, fill="#1E293B")
            yy += 48

    def arrow(x1, y1, x2, y2, color="#2E74B5"):
        draw.line((x1, y1, x2, y2), fill=color, width=8)
        draw.polygon([(x2, y2), (x2 - 20, y2 - 13), (x2 - 20, y2 + 13)], fill=color)

    arrow(320, 360, 380, 360)
    arrow(660, 360, 720, 360)
    arrow(1120, 360, 1180, 360)
    arrow(1480, 360, 1540, 360)

    rounded_rectangle(draw, (190, 690, 1610, 930), 24, "#F8FAFC", "#D7DEE8", 3)
    draw.text((235, 720), "可审计评测闭环", font=box_title, fill="#17365D")
    stages = [
        (250, "冻结金标集"),
        (530, "离线回放"),
        (810, "轨迹评分"),
        (1090, "人工纠错"),
        (1370, "回归测试"),
    ]
    for idx, (x, label) in enumerate(stages):
        rounded_rectangle(draw, (x, 800, x + 190, 875), 18, "#FFFFFF", "#2E74B5", 3)
        tw = draw.textbbox((0, 0), label, font=small_font)[2]
        draw.text((x + (190 - tw) / 2, 820), label, font=small_font, fill="#1E293B")
        if idx < len(stages) - 1:
            arrow(x + 190, 838, stages[idx + 1][0] - 15, 838, "#667085")

    draw.text((60, 980), "所有结论可回放到：原始来源 → 声明 → 证据片段 → 工具调用 → 模型/提示词版本 → 人工覆盖记录", font=body_font, fill="#667085")
    image.save(path, quality=95)


def add_figure(doc, image_path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=2, after=4, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run()
    shape = run.add_picture(str(image_path), width=Inches(6.45))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", "财经事件雷达V3.0总体架构")
    cp = doc.add_paragraph()
    set_paragraph_spacing(cp, before=0, after=10, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = cp.add_run(caption)
    set_run_font(r, size=9, italic=True, color=MUTED)


def add_reference(doc, index, title, url=None):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=4, line=1.15, align=WD_ALIGN_PARAGRAPH.LEFT)
    r = p.add_run(f"[{index}] {title}：")
    set_run_font(r, size=9.5, color=INK)
    if url:
        add_hyperlink(p, "查看来源", url)


def build():
    make_architecture_diagram(ARCHITECTURE_IMAGE)
    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    bullet_id, decimal_id = add_num_defs(doc)
    props = doc.core_properties
    props.title = "基于多源证据链与时序真值校验的金融事件情报 Agent"
    props.subject = "北京林业大学理学院2026实训高难度自主选题项目计划书"
    props.author = "Finance Radar Project Team"
    props.keywords = "Finance Radar, Evidence Agent, Financial Event Intelligence, LangGraph, Audit"

    # Cover — proposal_centerpiece.
    add_para(doc, "HIGH-DIFFICULTY PROJECT PROPOSAL  ·  V3.0", align=WD_ALIGN_PARAGRAPH.CENTER,
             before=34, after=16, line=1.0, size=10.5, color=BLUE)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=9, line=1.05, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("基于多源证据链与时序真值校验的")
    set_run_font(r, size=22, bold=True, color=NAVY)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=10, line=1.05, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("金融事件情报 Agent")
    set_run_font(r, size=30, bold=True, color=NAVY)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=8, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("声明抽取 · 权威证据检索 · 冲突验证 · 时序身份 · 可审计回放")
    set_run_font(r, size=13.5, color=DARK_BLUE)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=24, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("自主高难度创新选题申报稿 / 实训交付蓝图")
    set_run_font(r, size=11, italic=True, color=MUTED)

    cover = doc.add_table(rows=3, cols=2)
    values = [
        ("版本", "V3.0", "编制日期", "2026年7月17日"),
        ("当前状态", "可运行研究原型", "申报等级", "竞赛/高难度创新级"),
        ("项目边界", "研究与监控，不构成投资建议", "交易权限", "无；行情全部只读"),
    ]
    for row_idx, values_row in enumerate(values):
        for col_idx in range(2):
            label = values_row[col_idx * 2]
            value = values_row[col_idx * 2 + 1]
            cell = cover.rows[row_idx].cells[col_idx]
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, before=2, after=2, line=1.1, align=WD_ALIGN_PARAGRAPH.CENTER)
            r1 = p.add_run(f"{label}\n")
            set_run_font(r1, size=8.5, bold=True, color=MUTED)
            r2 = p.add_run(value)
            set_run_font(r2, size=10.2, bold=True, color=NAVY)
            shade_cell(cell, LIGHT)
    set_repeat_table_header(cover.rows[0])
    set_table_geometry(cover, [4680, 4680])

    add_para(doc, "", before=0, after=18, line=1.0)
    add_callout(
        doc,
        "一句话价值",
        "本项目不预测涨跌，而是证明一条金融消息在什么时间、由哪份一手文件支持、处于哪个法律或业务阶段、影响哪些资产，并在证据不足时明确拒绝下结论。",
        fill=LIGHT_BLUE,
    )
    add_para(doc, "项目严格区分“已实现能力”和“实训期计划能力”，不以计划冒充成品，不以市场波动反向证明消息为真。",
             align=WD_ALIGN_PARAGRAPH.CENTER, before=14, after=0, line=1.0, size=9.5, color=MUTED)
    doc.add_page_break()

    add_heading(doc, "执行摘要", 1)
    add_para(doc, "财经新闻的难点不是“看见”，而是“证明”。同一事件会被多次转载、不同来源互相矛盾，证券代码会因重组或退市变化，破产申请、重组计划获批、计划生效和旧普通股注销又是不同法律阶段。普通聚合器往往把文章当作事件，把模型语气当作置信度，最终形成不可复核的结论。")
    add_para(doc, "V3.0把现有财经事件雷达从多源采集与规则研究系统，升级为证据优先的金融事件情报 Agent。Agent将原始文本拆成原子声明，规划需要查找的一手证据，调用SEC、监管机构、公司公告等只读工具，生成声明—证据矩阵，并对每条声明标记支持、反驳或证据不足。最终等级由确定性安全门和人工复核控制，模型不能自动晋升S级。")
    add_callout(doc, "申报判断", "项目应以自主高难度创新选题申报，而不是作为普通A10“多源新闻聚合Agent”。现有系统已经具备真实数据底座；实训期的核心任务是补齐运行时AI、可视化演示、可复现评测和规范化过程证据。", fill=LIGHT_GREEN, accent=GREEN)

    add_heading(doc, "V3.0的三项决定", 2)
    add_list(doc, "把核心对象从“文章”改为可版本化的Canonical Event，并以声明—证据边表达结论来源。", decimal_id)
    add_list(doc, "使用一个有边界、可恢复的Evidence Agent状态图，不堆砌多个角色Agent，也不添加装饰性聊天框。", decimal_id)
    add_list(doc, "用冻结金标集、拒绝对照、消融实验和轨迹评分证明系统有效，而不是只展示页面和个别成功案例。", decimal_id)

    add_heading(doc, "当前基线与实训目标", 2)
    add_table(
        doc,
        ["层次", "当前事实", "V3.0实训目标", "状态"],
        [
            ("数据底座", "18个注册来源；1,160个规范事件；3,556条原始观测", "保持只读采集稳定，新增Agent轨迹与声明表", "已实现 + 计划增强"),
            ("证据与版本", "2,386条证据；2,081个事件版本", "重大声明绑定精确证据片段和快照哈希", "已实现 + 计划增强"),
            ("裁决样本", "792个历史裁决：441 verified / 351 rejected", "冻结时间切分测试集，建立三方案对照实验", "已实现 + 计划增强"),
            ("AI运行时", "当前主要为规则、检索与人工裁决", "声明抽取、检索规划、证据验证与带引用摘要", "计划实现"),
            ("演示层", "报告、命令行和Telegram outbox", "只读Streamlit事件工作台与轨迹回放", "计划实现"),
            ("工程过程", "232项测试通过；审计PASS", "有效Git、3个Sprint、.agent文档、核心覆盖率≥80%", "计划补齐"),
        ],
        [1260, 3000, 3600, 1500],
        font_size=8.8,
        center_cols=(3,),
        status_col=3,
    )
    doc.add_page_break()

    add_heading(doc, "1. 项目定位与差异化", 1)
    add_heading(doc, "1.1 问题定义", 2)
    add_para(doc, "目标用户是需要快速研究美股、ETF、宏观、商品、外汇和主要加密资产事件的个人研究者。用户真正需要的不是更多标题，而是一个可以回答五个问题的系统：发生了什么、涉及谁、证据在哪里、来源是否冲突、当前为什么能或不能确认。")
    add_para(doc, "项目选择“可信事件形成”作为核心研究问题：在多源、冲突、时序变化的金融信息环境中，如何让Agent的每一个重大结论都可证明、可质疑、可回放，并在证据不足时可靠地拒绝下结论。")

    add_heading(doc, "1.2 与标准A10项目的本质区别", 2)
    add_table(
        doc,
        ["维度", "标准多源新闻聚合Agent", "财经事件雷达V3.0"],
        [
            ("处理对象", "新闻文章与摘要", "跨来源、跨阶段、可版本化的金融事件"),
            ("AI角色", "分类、摘要、推荐、追问", "原子声明抽取、证据研究规划、冲突验证、带引用摘要"),
            ("可信依据", "模型输出或新闻数量", "权威来源、来源独立组、精确证据片段、确定性安全门"),
            ("时间语义", "通常使用文章发布时间", "事件发生时点、来源发布时间、接收时间、确认时间分离"),
            ("资产语义", "文章级统一情绪", "事件事实与逐资产影响分离，允许不同资产方向相反"),
            ("评价方式", "摘要质量与用户体验", "事实精度、引用完整率、拒绝对照误报、身份准确率、轨迹合规"),
            ("安全边界", "通常未定义", "无交易工具；行情仅事后审计；证据不足即拒绝"),
        ],
        [1380, 3340, 4640],
        font_size=9.1,
    )

    add_heading(doc, "1.3 核心创新", 2)
    for text in [
        "从“新闻摘要”升级为“声明—证据—结论”：每条重大结论必须能定位到支持或反驳它的原文片段。",
        "发现源与证明源分离：Telegram、X和聚合器负责发现，SEC、法院、监管机构和发行人原文负责确认。",
        "时点证券身份解析：在事件发生时识别当时的公司、证券、交易场所和代码，避免把后来的Q后缀或新证券身份倒灌到过去。",
        "显式表示INSUFFICIENT：证据不足是一个可解释的正式结果，而不是模型失败或被迫猜测。",
        "评测驱动的研究闭环：把人工纠错、拒绝对照和失败轨迹回灌成冻结数据集与回归测试。",
    ]:
        add_list(doc, text, bullet_id)

    add_heading(doc, "2. 现有基础与真实性证据", 1)
    add_heading(doc, "2.1 已经运行的底座", 2)
    add_table(
        doc,
        ["证据", "当前值", "说明"],
        [
            ("数据库", "SQLite Schema 12", "原始观测、事件、版本、证据、任务、行情与审计对象分离"),
            ("来源", "18个注册来源", "包含Fed、BLS、SEC、CFTC、FDA、FTC、FDIC、OpenNews及只读行情"),
            ("数据规模", "1,160事件 / 3,556原始观测", "不是空原型，已形成可回放样本"),
            ("证据规模", "2,386证据 / 2,081版本", "支持事件随官方材料变化而更新"),
            ("人工裁决", "441验证 / 351拒绝", "拒绝样本用于度量误报，不作为无用数据删除"),
            ("测试", "232项通过", "覆盖事件发现、SEC证据、裁决、事件链、只读行情和通知幂等"),
            ("安全审计", "PASS；19类违规计数为0", "无交易、无自动晋级、候选与确认隔离"),
        ],
        [1700, 2200, 5460],
        font_size=9.2,
        center_cols=(1,),
    )
    add_callout(doc, "真实性声明", "以上数字来自2026年7月16日至17日的本地数据库、审计报告和测试复跑。运行时AI、可视化Dashboard、覆盖率报告和有效Git过程尚未完成，均在本文中明确标记为计划目标。", fill=LIGHT_GOLD, accent=GOLD)

    add_heading(doc, "2.2 已证明有价值的困难案例", 2)
    add_para(doc, "项目最有说服力的不是数据量，而是对“看起来相似、法律结果不同”的案例进行区分。NINEQ初始8-K只描述拟注销旧股，直到103天后的10-Q证明计划已生效、旧普通股无对价注销，事件才允许从A++进入S。WOLF2则证明旧股被取消但持有人获得新股恢复时不能升S。")
    add_para(doc, "这组正反案例展示了三项不可由普通新闻摘要替代的能力：法律阶段识别、事件时间版本化、旧证券与新证券的身份连续性判断。答辩时将使用同一界面回放两个案例，让评委看到系统不仅会“确认”，还会正确地“拒绝确认”。")
    doc.add_page_break()

    add_heading(doc, "3. V3.0总体架构", 1)
    add_figure(doc, ARCHITECTURE_IMAGE, "图1  证据驱动、时序感知、可审计的金融事件情报Agent架构")
    add_heading(doc, "3.1 数据面与控制面", 2)
    add_table(
        doc,
        ["平面", "职责", "关键对象", "故障策略"],
        [
            ("数据面", "采集、声明抽取、证据检索、验证、事件版本与输出", "RawObservation、EventClaim、EvidenceEdge、CanonicalEvent", "先落盘、幂等重试、失败不丢原文"),
            ("控制面", "来源配置、配额、模型/提示词版本、轨迹、人工中断与回放", "SourceRegistry、AgentDecision、Trace、ReviewTask", "检查点恢复、预算上限、人工接管"),
            ("审计面", "离线金标评测、线上失败监测、回归与消融", "GoldSet、EvalRun、GuardrailResult", "失败样本只追加，不静默覆盖"),
        ],
        [1260, 2800, 2860, 2440],
        font_size=9.0,
    )

    add_heading(doc, "3.2 单Agent状态图", 2)
    add_para(doc, "系统采用一个有界Evidence Agent，而不是多角色对话。LangGraph只负责显式状态、条件路由、检查点和人工中断；每个节点只完成一件事，输入输出均可审计。")
    states = [
        "DISCOVERED：发现源进入，尚不代表事件真实。",
        "EXTRACTED：已得到结构化原子声明。",
        "NEEDS_EVIDENCE：Agent生成官方检索计划并调用白名单工具。",
        "SUPPORTED / CONTRADICTED / INSUFFICIENT：逐声明形成证据关系。",
        "HUMAN_REVIEW：达到重大性阈值或存在冲突时暂停等待人工复核。",
        "PUBLISHED：通过确定性安全门后发布事件卡；后续证据生成新版本。",
    ]
    for text in states:
        add_list(doc, text, bullet_id)

    add_heading(doc, "3.3 模型与代码的职责边界", 2)
    add_table(
        doc,
        ["能力", "LLM/语义模型", "确定性代码", "人工"],
        [
            ("声明抽取", "生成结构化候选声明", "JSON Schema验证、字段范围检查", "抽样纠错"),
            ("证据检索", "规划查询、语义重排", "域名白名单、限流、缓存、循环预算", "补充特殊来源"),
            ("证据判断", "提出支持/反驳/不足", "时间、数字、身份、最终性规则复核", "重大结论审批"),
            ("摘要", "生成带引用的人读说明", "引用覆盖、禁止无来源声明", "必要时编辑"),
            ("等级", "只能提出建议", "最终安全门；禁止自动S；无交易路径", "最终裁决与覆盖记录"),
        ],
        [1320, 2740, 3220, 2080],
        font_size=9.0,
    )

    add_heading(doc, "4. Evidence Agent核心设计", 1)
    add_heading(doc, "4.1 三个模型节点", 2)
    add_list(doc, "声明抽取器：把文本拆成主体、动作、对象、时间、金额、事件类型、重大性和不确定原因；输出EventClaim。", decimal_id)
    add_list(doc, "证据研究规划器：根据事件类型决定应查SEC表单、监管公告、法院文件还是公司IR，并生成最多三轮的有预算查询。", decimal_id)
    add_list(doc, "证据验证与摘要器：逐条标记supports、contradicts或contextual，生成附精确引用的结论，并解释尚缺什么证据。", decimal_id)

    add_heading(doc, "4.2 严格结构化输出", 2)
    add_table(
        doc,
        ["对象", "核心字段", "用途"],
        [
            ("EventClaim", "claim_id、subject、predicate、object、event_type、event_time、instrument_identity_at_event、status", "把自然语言拆成可逐条核验的原子声明"),
            ("EvidenceEdge", "source_url、snapshot_hash、authority_tier、exact_excerpt、offset、relation、claim_id", "证明具体哪段材料支持或反驳哪条声明"),
            ("AgentDecision", "trace_id、model_snapshot、prompt_version、tool_calls、evidence_ids、guardrails、latency、human_override", "实现模型、工具、提示词与人工覆盖的完整回放"),
        ],
        [1560, 4680, 3120],
        font_size=8.9,
    )
    add_callout(doc, "Fail closed", "模型输出不符合Schema、来源中断、查询预算耗尽、存在未解决冲突或缺少最终性文件时，系统必须进入INSUFFICIENT或HUMAN_REVIEW，不能猜测性发布。", fill=LIGHT_RED, accent=RED)

    add_heading(doc, "4.3 两级证据检索", 2)
    add_para(doc, "第一阶段使用CIK、证券代码、公司名、事件类型、时间窗口、表单类型和BM25等高召回条件搜索；第二阶段使用语义模型对候选证据重排，并检查段落是否含有可支持声明的具体事实。搜索结果摘要不能直接作为证明，只有原始页面或文件中的精确片段可以形成EvidenceEdge。")
    add_para(doc, "对SEC自动访问设置可识别User-Agent、每秒不超过10次的总限流、ETag/Last-Modified缓存、指数退避与熔断。对新闻、Telegram和网页内容一律视为不可信输入，清除脚本和隐藏文本，并阻止其成为高优先级提示词或改变工具权限。")
    doc.add_page_break()

    add_heading(doc, "5. AI禁飞区与学生亲手实现", 1)
    add_para(doc, "为满足课程“每个项目3个核心功能必须手写并逐行解释”的要求，以下模块被定义为AI禁飞区。AI可以提供原理资料和测试想法，但不得直接生成最终实现；提交前保留设计草稿、手写代码记录、单元测试和代码走查提纲。")
    add_table(
        doc,
        ["禁飞区", "学生必须实现的核心", "现场可验证方式", "失败后果"],
        [
            ("FZ-1 Agent循环与事件聚类", "状态路由、事件指纹、同源修订、跨来源合并、事件链primary/consequence边界", "注入重复标题、同日不同事件和跨家族后果，现场解释路由", "重复爆炸或错误合并"),
            ("FZ-2 证据检索与忠实度校验", "查询规划预算、精确片段绑定、声明支持率、引用完整率、矛盾检测", "删去关键证据或加入反驳证据，观察结论降级", "模型摘要无依据"),
            ("FZ-3 时序身份与最终评级安全门", "事件时点证券身份、法律阶段、来源等级、禁止自动S、无交易工具", "把破产申请与计划生效混淆，或把新证券身份倒灌到旧事件", "高风险错误确认"),
        ],
        [1450, 3650, 2900, 1360],
        font_size=8.7,
    )
    add_heading(doc, "5.1 需要真正理解的核心循环", 2)
    add_callout(doc, "伪代码边界", "persist raw → extract claims → resolve event-time identity → search primary evidence → build claim-evidence matrix → deterministic checks → interrupt for human review → publish version。每一步都必须幂等、可回放，且在任何失败点都不能绕过安全门。", fill=LIGHT_BLUE)

    add_heading(doc, "5.2 Bug注入与即兴修改准备", 2)
    add_table(
        doc,
        ["考核场景", "预置问题", "30分钟目标"],
        [
            ("Bug注入1", "同一来源编辑稿被误判成独立确认源", "修复independence_group并补回归测试"),
            ("Bug注入2", "SEC时间字段时区错误导致事件版本乱序", "修复时间归一化并解释t0选择"),
            ("Bug注入3", "模型把反驳片段当成支持片段", "修复关系验证并展示降级到HUMAN_REVIEW"),
            ("即兴修改", "新增“只显示缺失证据”的审核筛选器", "修改查询/API/UI并通过测试"),
        ],
        [1500, 4300, 3560],
        font_size=9.0,
    )

    add_heading(doc, "6. 数据模型、来源治理与安全边界", 1)
    add_heading(doc, "6.1 来源等级与独立性", 2)
    add_table(
        doc,
        ["等级", "来源", "允许作用", "禁止作用"],
        [
            ("P0", "监管、法院、统计机构、交易所正式文件", "确认文件明确陈述的事实", "超出原文推断责任、因果或收益"),
            ("P1", "公司IR、正式新闻稿、官方讲话", "确认主体自身声明", "替代监管或法院的最终认定"),
            ("P2", "信誉媒体、OpenNews等聚合API", "发现、上下文、候选事件", "单独确认重大事件"),
            ("P3", "社交帖子、匿名爆料、二次转述", "产生线索", "主动标记S/A已确认"),
        ],
        [900, 2700, 2880, 2880],
        font_size=9.0,
        center_cols=(0,),
    )
    add_para(doc, "同一Reuters原文的媒体转载、X转发和聚合器记录只能属于一个independence_group，不能被计算为三个独立确认源。原始正文默认不长期全文分发，只保存必要片段、哈希、元数据和原始链接。")

    add_heading(doc, "6.2 事件事实、资产影响与市场观察分离", 2)
    add_table(
        doc,
        ["对象", "回答的问题", "不能替代什么"],
        [
            ("Canonical Event", "发生了什么、处于什么阶段、是否已被证据支持", "不能直接给出投资建议"),
            ("Asset Impact", "事件与某个具体资产的关系、方向候选、窗口和不确定性", "不能决定事件事实真假"),
            ("Market Observation", "1/5/15/60分钟及收盘后的价格、异常收益和成交量", "不能反向证明新闻为真"),
            ("Provider Assessment", "外部聚合器或模型给出的分数与摘要", "不能覆盖本地证据裁决"),
        ],
        [1900, 3900, 3560],
        font_size=9.1,
    )

    add_heading(doc, "6.3 不可逾越的安全边界", 2)
    for text in [
        "系统没有下单、撤单、持仓、账户和资金工具；IBKR仅只读，Binance服务器只调用公共行情端点。",
        "模型节点看不到密钥，不允许自由文本直接构造任意URL、SQL或系统命令。",
        "所有外部文本均视为潜在提示词注入；工具使用域名白名单、次数预算、超时和只读权限。",
        "市场数据只能在事件形成后做影响审计，不能进入事实确认或训练标签形成过程。",
        "重大事件没有P0/P1证据或独立可靠来源时不得自动确认；S级重大声明支持率必须为100%。",
    ]:
        add_list(doc, text, bullet_id)
    doc.add_page_break()

    add_heading(doc, "7. 数据集、模型与可复现评测", 1)
    add_heading(doc, "7.1 金标集设计", 2)
    add_para(doc, "从现有792个历史裁决中冻结一组不参与提示词调试的测试集。切分优先按时间、发行人和事件链分组，避免同一公司、同一事件的近重复文本同时出现在训练与测试中。每条样本只允许使用事件时点前可获得的信息，事后行情结果不得成为输入特征。")
    add_table(
        doc,
        ["样本族", "必须覆盖的边界"],
        [
            ("真实事件", "破产申请、计划生效、旧股注销、重大融资、强制退市、监管暂停"),
            ("拒绝对照", "现金并购退市、SPAC单位终止、正常反向拆股、假设性清算文本、旧事件重复披露"),
            ("时序/身份", "Q后缀、OTC转移、代码变更、新旧证券、同一事件后续版本"),
            ("系统异常", "权威源中断、无证据、冲突证据、恶意提示词、Schema错误"),
        ],
        [2100, 7260],
        font_size=9.3,
    )

    add_heading(doc, "7.2 不训练大模型，先训练窄任务", 2)
    add_para(doc, "V3.0不宣称训练专用金融大模型。通用模型负责结构化语义任务；现有裁决数据只用于一个可解释的小型候选筛选器或证据相关性重排器。建议比较TF-IDF+逻辑回归、现有规则和LLM few-shot三种基线，输出“是否值得进入人工审核”的优先级，不直接决定S/A等级。")
    add_para(doc, "主要指标采用Precision、Recall、F1、PR-AUC、Recall@K与校准误差；对拒绝对照单独报告误晋升率。模型版本、特征版本、训练切分和随机种子全部写入实验记录。")

    add_heading(doc, "7.3 四层事实忠实度", 2)
    add_list(doc, "确定性检查：实体、时间、数字、证券身份和最终性条款。", decimal_id)
    add_list(doc, "原子事实支持率：摘要拆成原子声明，计算被权威证据支持的比例。", decimal_id)
    add_list(doc, "引用正确性与完整性：相邻引用是否真的支持声明，所有重大声明是否都有引用。", decimal_id)
    add_list(doc, "QA与蕴含复核：从摘要生成问题，比较来源答案与摘要答案，并检查矛盾。", decimal_id)

    add_heading(doc, "7.4 验收指标", 2)
    add_table(
        doc,
        ["能力", "验收目标", "测量方法"],
        [
            ("结构化输出合规率", "100%", "Schema验证；失败即拒绝"),
            ("权威证据Recall@10", "≥90%", "冻结查询集"),
            ("实体/事件/时间抽取Macro-F1", "≥90%", "人工标注测试集"),
            ("事件去重Pairwise-F1", "≥90%", "同事件/非同事件对"),
            ("引用正确率与重大声明完整率", "均≥95%", "人工双审+自动检查"),
            ("已确认摘要原子事实支持率", "≥95%", "声明—证据矩阵"),
            ("S级重大声明支持率", "100%", "硬性安全门"),
            ("拒绝对照错误晋升S", "0", "351个拒绝对照及新增对抗集"),
            ("时点证券身份准确率", "≥99%", "身份金标集"),
            ("审计轨迹覆盖率", "100%", "trace_id完整性审计"),
            ("未授权/交易工具调用", "0", "工具白名单与审计"),
            ("核心模块测试覆盖率", "≥80%", "coverage报告"),
        ],
        [2850, 1650, 4860],
        font_size=8.9,
        center_cols=(1,),
    )

    add_heading(doc, "7.5 对照与消融实验", 2)
    add_table(
        doc,
        ["实验组", "目的"],
        [
            ("现有规则系统", "提供当前可复现基线"),
            ("纯LLM系统", "展示无证据约束时的幻觉、冲突和身份错配风险"),
            ("证据驱动混合系统", "V3.0最终方案"),
            ("去掉权威源优先级", "量化来源治理的贡献"),
            ("去掉忠实度复核", "量化引用和声明校验的贡献"),
            ("去掉时序身份解析", "量化代码变化与法律阶段错配的影响"),
        ],
        [3200, 6160],
        font_size=9.2,
    )
    doc.add_page_break()

    add_heading(doc, "8. 出彩的产品演示设计", 1)
    add_heading(doc, "8.1 界面不是聊天框", 2)
    add_para(doc, "演示层采用只读Streamlit工作台，围绕证据和状态变化组织，不把聊天框作为主界面。首页只显示真实事件流、来源健康和审核队列；事件详情页展示声明、证据、冲突、事件时间线、Agent轨迹和市场观察。")
    add_table(
        doc,
        ["页面/组件", "用户能回答的问题"],
        [
            ("实时事件流", "现在发生了什么？哪些仍待核验？"),
            ("声明—证据矩阵", "为什么确认？哪条证据支持或反驳？"),
            ("证据缺口", "为什么还不能确认？下一步应查什么？"),
            ("事件版本时间线", "结论何时、因哪份文件发生变化？"),
            ("Agent轨迹回放", "调用了哪些工具、经历了哪些安全门和人工覆盖？"),
            ("反事实切换", "如果去掉某条证据，评级是否降级？"),
            ("市场观察", "事件后价格如何变化？明确标记为审计而非事实证明"),
        ],
        [2500, 6860],
        font_size=9.3,
    )

    add_heading(doc, "8.2 三分钟答辩主线", 2)
    add_table(
        doc,
        ["时间", "演示动作", "评委看到的能力"],
        [
            ("0:00–0:25", "聚合新闻或Telegram出现一条高影响传闻", "实时发现；明确标记未核验"),
            ("0:25–0:55", "Agent拆出原子声明并生成证据计划", "AI深度进入核心流程"),
            ("0:55–1:30", "自动查询SEC/监管原文，展示支持与反驳片段", "工具调用、引用与冲突处理"),
            ("1:30–2:00", "因缺少最终性文件拒绝升S，人工审核暂停", "可拒绝、不盲信模型"),
            ("2:00–2:30", "回放后续官方文件到达，事件产生新版本", "时序真值与版本化"),
            ("2:30–2:50", "对照NINEQ与WOLF2", "真阳性与困难负样本"),
            ("2:50–3:00", "注入“忽略规则并下单”的恶意文本", "提示词注入防护与无交易工具"),
        ],
        [1200, 4100, 4060],
        font_size=9.0,
        center_cols=(0,),
    )
    add_callout(doc, "答辩金句", "我们没有让模型更自信，而是让系统更会证明、更会拒绝、更能回放。", fill=LIGHT_GREEN, accent=GREEN)

    add_heading(doc, "9. 12天、3个Sprint实施计划", 1)
    add_table(
        doc,
        ["阶段", "时间", "主要工作", "可验收出口"],
        [
            ("准备", "Day 1", "完成高难度选题评审；建立有效Git；填写团队角色；冻结安全边界", "基线提交、角色表、选题评审记录"),
            ("设计", "Day 2–3", "人工完成架构、接口、数据模型、测试、部署设计；冻结首版金标集", "5份人工设计产物 + .agent目录"),
            ("Sprint 1", "Day 4–6", "Evidence Agent状态图、EventClaim/EvidenceEdge契约、单案例垂直闭环", "NINEQ回放可运行；Schema测试通过"),
            ("Sprint 2", "Day 7–9", "两级检索、忠实度评估、窄任务基线、拒绝对照与Dashboard", "三方案对照报告；工作台可演示"),
            ("Sprint 3", "Day 10–11", "安全门、提示词注入测试、覆盖率、Bug注入、部署与演练", "覆盖率≥80%；审计PASS；一键启动"),
            ("答辩", "Day 12", "三分钟主线、代码走查、即兴修改、个人贡献说明", "演示录像/现场演示 + 答辩包"),
        ],
        [1200, 1050, 4550, 2560],
        font_size=8.9,
        center_cols=(1,),
    )

    add_heading(doc, "9.1 Sprint验收原则", 2)
    for text in [
        "每个用户故事都有验收标准；未通过测试或没有证据的功能不计完成。",
        "每个功能一个可解释提交；单次改动超过300行必须说明原因并拆分评审。",
        "每个Sprint记录AI使用、采纳/拒绝理由、代码审查和复盘，不补写虚假过程。",
        "现有原型作为baseline commit导入；历史开发过程不伪造，实训期过程从Day 1真实记录。",
    ]:
        add_list(doc, text, bullet_id)

    add_heading(doc, "10. 团队、AI协作与工程规范", 1)
    add_heading(doc, "10.1 团队角色", 2)
    add_table(
        doc,
        ["角色", "职责", "建议产出"],
        [
            ("PO / 架构负责人", "需求优先级、事件模型、验收标准、答辩叙事", "用户故事、架构决策、主流程代码"),
            ("SM / 后端负责人", "Sprint节奏、采集与事件账本、部署", "Backlog、接口与运行脚本"),
            ("QA / 评测负责人", "独立设计金标、对照实验、Bug注入、覆盖率", "测试设计、评测报告、缺陷记录"),
            ("前端 / 演示负责人", "只读工作台、轨迹回放、三分钟演示", "Streamlit页面、演示脚本"),
        ],
        [1900, 4100, 3360],
        font_size=9.1,
    )
    add_para(doc, "实际成员为2–4人时可合并角色，但QA必须保持独立评测视角；所有成员都必须提交代码，并在最终材料中按提交、评审、测试和文档证据说明个人贡献。")

    add_heading(doc, "10.2 .agent项目记忆", 2)
    add_table(
        doc,
        ["文件", "必须记录的内容"],
        [
            (".agent/architecture.md", "系统边界、状态图、组件职责和关键数据流"),
            (".agent/coding_conventions.md", "Python规范、命名、异常、日志、幂等和测试约定"),
            (".agent/api_contracts.md", "EventClaim、EvidenceEdge、AgentDecision和只读工具契约"),
            (".agent/decisions.md", "关键取舍、备选方案、日期、负责人和影响"),
            (".agent/fixes.md", "Bug、根因、修复、回归测试和是否由AI建议"),
            (".agent/forbidden_zones.md", "三处禁飞区、负责人、代码路径、走查问题"),
        ],
        [2850, 6510],
        font_size=9.2,
    )

    add_heading(doc, "10.3 五份必须由人主导的设计产物", 2)
    add_list(doc, "架构设计：为什么采用单Agent状态图，为什么不堆多Agent。", decimal_id)
    add_list(doc, "接口设计：模型、工具、数据库和UI之间的严格契约。", decimal_id)
    add_list(doc, "数据模型：声明、证据、事件、版本、身份和轨迹关系。", decimal_id)
    add_list(doc, "测试设计：金标切分、边界案例、对抗样本、消融和覆盖率。", decimal_id)
    add_list(doc, "部署设计：一键启动、配置、密钥、备份、恢复和只读权限。", decimal_id)

    add_heading(doc, "11. 测试、部署与运行可靠性", 1)
    add_heading(doc, "11.1 测试金字塔", 2)
    add_table(
        doc,
        ["层级", "范围", "目标"],
        [
            ("单元测试", "解析、聚类、时间、身份、安全门、Schema", "核心模块覆盖率≥80%"),
            ("契约测试", "SEC/OpenNews/行情/Telegram适配器与结构化模型输出", "离线fixture可复跑"),
            ("集成测试", "原始落盘→Agent→人工中断→事件版本→输出", "至少3条完整垂直链路"),
            ("评测测试", "金标、拒绝对照、冲突、无证据和注入攻击", "关键指标达到第7.4节目标"),
            ("恢复测试", "进程中断、断线、429、超时、重复消息", "不丢数据、不重复发布"),
        ],
        [1800, 4700, 2860],
        font_size=9.1,
    )

    add_heading(doc, "11.2 一键演示与部署", 2)
    add_para(doc, "实训交付采用Windows本地运行优先：一个启动脚本初始化SQLite、检查环境、启动只读API与Streamlit，并提供离线replay模式。网络或外部API不可用时，仍可用冻结快照完整演示Agent轨迹和评测。Docker作为可选增强，不作为演示唯一依赖。")
    add_para(doc, "密钥仅存放于.env，提交.env.example；报告、日志、数据库正文和截图不得出现密钥。数据库每日备份，原始观测不可变，模型和提示词升级必须保留版本，所有人工覆盖保留操作者、时间和理由。")

    add_heading(doc, "12. 风险登记与降级策略", 1)
    add_table(
        doc,
        ["风险", "影响", "降级与验收"],
        [
            ("模型幻觉或Schema漂移", "产生无依据声明", "严格Schema；引用覆盖校验；失败进入INSUFFICIENT"),
            ("提示词注入", "诱导越权工具调用", "不可信文本隔离；工具白名单；注入测试；无交易工具"),
            ("SEC限流或中断", "确认延迟", "10 req/s以下；缓存、退避、熔断；离线fixture"),
            ("聚合器不可用", "发现覆盖下降", "P2可关闭；P0/P1官方链路继续运行"),
            ("样本泄漏", "离线指标虚高", "按时间/发行人/事件链分组切分；测试集冻结"),
            ("人工审核成为瓶颈", "候选积压", "窄任务优先级、证据缺口队列、每族吞吐指标"),
            ("12天范围过大", "完不成展示", "优先NINEQ/WOLF2闭环；训练、Docker和更多来源为可选项"),
            ("Git历史缺失", "过程分丢失", "Day 1建立真实baseline，后续不伪造历史、细粒度提交"),
        ],
        [2000, 3000, 4360],
        font_size=9.0,
    )
    doc.add_page_break()

    add_heading(doc, "13. 交付物与最终验收", 1)
    add_table(
        doc,
        ["类别", "交付物", "验收证据"],
        [
            ("运行系统", "官方源采集、Evidence Agent、确定性安全门、事件账本、Streamlit、Telegram", "一键启动与离线replay"),
            ("数据与模型", "EventClaim/EvidenceEdge/AgentDecision；冻结金标；窄任务基线", "数据卡、切分清单、模型卡、指标报告"),
            ("工程", "有效Git、CI或本地质量门、测试、coverage、审计", "提交记录、232+测试、覆盖率≥80%、审计PASS"),
            ("过程", "3 Sprint、Backlog、每日站会、AI日志、评审记录", "过程包与个人贡献矩阵"),
            ("人主导设计", "架构、接口、数据模型、测试、部署", "5份签名/负责人明确的设计文档"),
            ("答辩", "三分钟主线、禁飞区走查、Bug注入、即兴修改", "演示脚本、问题库、现场通过"),
        ],
        [1450, 4550, 3360],
        font_size=9.0,
    )

    add_heading(doc, "13.1 Definition of Done", 2)
    for text in [
        "功能有用户故事、验收标准、代码、测试、文档和可重复演示。",
        "AI输出有Schema、证据引用、模型/提示词版本和完整trace_id。",
        "任何重大结论都能回答“为什么确认、为什么不能确认、什么证据会改变结论”。",
        "任何外部中断或模型失败都不能启用交易、自动升S或丢失原始观测。",
        "每位成员能逐行解释本人禁飞区代码，并在30分钟内完成限定修改。",
    ]:
        add_list(doc, text, bullet_id)

    add_heading(doc, "14. 校方高分标准映射", 1)
    add_para(doc, "项目不以“功能很多”作为高分依据，而把每个评分点绑定到可查验产物。过程评价约占60%，因此Git、Sprint、AI日志、代码走查和团队协作与最终功能同等重要。")
    add_table(
        doc,
        ["评分维度", "V3.0得分抓手", "现场证据"],
        [
            ("工程实践", "真实外部源、持久化、恢复、只读部署、完整垂直闭环", "运行系统、审计、replay"),
            ("AI代码审查/协作", "AI日志、采纳/拒绝理由、轨迹评分、三处禁飞区", "Sprint记录、走查、个人解释"),
            ("工程规范", "有效Git、.agent、5份设计、接口契约、覆盖率", "仓库、文档、测试报告"),
            ("团队协作", "PO/SM/QA分工、Backlog、评审、个人贡献矩阵", "提交/评审/测试证据"),
            ("功能与创新", "声明—证据—结论、时序身份、冲突与拒绝、消融实验", "NINEQ/WOLF2主演示与指标"),
            ("答辩理解", "三分钟主线、逐行代码、Bug注入、即兴修改", "现场完成而非视频替代"),
        ],
        [1900, 4260, 3200],
        font_size=9.0,
    )
    add_callout(doc, "预期档位", "若完成运行时AI、Dashboard、真实Git/Sprint过程、禁飞区掌握和可复现评测，项目具备90–95+的现实竞争力；最高档仍以教师现场评审和团队个人表现为准。", fill=LIGHT_GREEN, accent=GREEN)

    add_heading(doc, "15. 立项结论", 1)
    add_para(doc, "Finance Radar最有价值的部分不是来源数量，也不是追求毫秒级宣传，而是已经形成了真实事件、证据版本和困难负样本。V3.0将这套底座重构为一个真正进入核心业务的Evidence Agent：模型负责理解与研究，确定性代码负责边界与安全，人负责重大结论，评测负责证明系统是否真的变好。")
    add_para(doc, "因此，本项目适合以“基于多源证据链与时序真值校验的金融事件情报Agent”申报自主高难度创新项目。最小成功不是做出一个漂亮新闻列表，而是完成一条任何人都能复核的链路：发现一条消息，拆出声明，找到或找不到权威证据，解释冲突，更新事件版本，并在证据不足时可靠地停下来。")
    add_callout(doc, "最终取舍", "优先做深一条可证明的事件闭环，再扩展来源和模型；优先做可复现评测，再追求更高分数；优先让评委看懂为什么系统拒绝，再展示它能生成什么。", fill=LIGHT_BLUE)

    add_heading(doc, "16. 主要依据与参考资料", 1)
    add_reference(doc, 1, "北京林业大学2026实训动员会（本地课程文件）")
    add_reference(doc, 2, "北京林业大学理学院2026年实现项目列表（本地课程文件）")
    add_reference(doc, 3, "LangGraph Overview and Persistence", "https://docs.langchain.com/oss/python/langgraph/overview")
    add_reference(doc, 4, "LangGraph Human-in-the-loop", "https://docs.langchain.com/oss/python/langchain/human-in-the-loop")
    add_reference(doc, 5, "SEC Developer Resources and Fair Access", "https://www.sec.gov/about/developer-resources")
    add_reference(doc, 6, "NIST AI 600-1 Generative AI Profile", "https://nvlpubs.nist.gov/nistpubs/ai/NIST.AI.600-1.pdf")
    add_reference(doc, 7, "RAGAS: Automated Evaluation of Retrieval Augmented Generation", "https://aclanthology.org/2024.eacl-demo.16/")
    add_reference(doc, 8, "FActScore: Fine-grained Atomic Evaluation of Factual Precision", "https://aclanthology.org/2023.emnlp-main.741/")
    add_reference(doc, 9, "ALCE: Enabling Large Language Models to Generate Text with Citations", "https://aclanthology.org/2023.emnlp-main.398/")
    add_reference(doc, 10, "W3C PROV-O: The PROV Ontology", "https://www.w3.org/TR/prov-o/")
    add_reference(doc, 11, "OpenAI practical guide to building agents", "https://openai.com/business/guides-and-resources/a-practical-guide-to-building-ai-agents/")
    add_reference(doc, 12, "OpenNews MCP", "https://github.com/6551team/opennews-mcp")
    add_para(doc, "注：第1、2项为本地课程文件名，未上传互联网；项目交付包中保留原文件。其余链接用于架构、风险治理与评测方法研究，不代表对相关服务形成强依赖。", size=9.2, color=MUTED, align=WD_ALIGN_PARAGRAPH.LEFT)

    add_para(doc, "— 文档结束 —", align=WD_ALIGN_PARAGRAPH.CENTER, before=24, after=0, line=1.0, size=9, color=MUTED)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
