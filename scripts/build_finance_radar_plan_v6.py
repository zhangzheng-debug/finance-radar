#!/usr/bin/env python3
"""Build the novice-friendly Finance Radar V6.0 internship proposal."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_finance_radar_plan_v3 import (
    BLUE,
    DARK_BLUE,
    GOLD,
    GREEN,
    LIGHT_BLUE,
    LIGHT_GOLD,
    LIGHT_GREEN,
    MUTED,
    NAVY,
    add_callout,
    add_heading,
    add_list,
    add_num_defs as _legacy_add_num_defs,
    add_page_number,
    add_para,
    add_reference,
    add_table,
    configure_styles,
    set_paragraph_spacing,
    set_run_font,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "financial_event_radar_project_proposal_v6_0_human.docx"


def add_num_defs(doc):
    """Create valid list definitions and force independent decimal lists to restart.

    The shared legacy helper appends new abstract numbering definitions after
    existing ``w:num`` nodes. Word tolerates many OOXML variations, but that
    ordering can make it merge visually identical lists and continue their
    numbering. Move the new definitions back into schema order and add an
    explicit start override for this document.
    """
    bullet_id, decimal_id = _legacy_add_num_defs(doc)
    numbering = doc.part.numbering_part.element

    abstract_by_id = {
        int(node.get(qn("w:abstractNumId"))): node
        for node in numbering.findall(qn("w:abstractNum"))
    }
    first_num = numbering.find(qn("w:num"))
    for num_id in (bullet_id, decimal_id):
        num = next(
            node
            for node in numbering.findall(qn("w:num"))
            if int(node.get(qn("w:numId"))) == num_id
        )
        abstract_ref = num.find(qn("w:abstractNumId"))
        abstract_id = int(abstract_ref.get(qn("w:val")))
        abstract = abstract_by_id[abstract_id]
        numbering.remove(abstract)
        numbering.insert(numbering.index(first_num), abstract)

    decimal_num = next(
        node
        for node in numbering.findall(qn("w:num"))
        if int(node.get(qn("w:numId"))) == decimal_id
    )
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    decimal_num.append(override)
    return bullet_id, decimal_id


def configure_document(doc: Document) -> None:
    """Resolve the narrative_proposal preset with a proposal_centerpiece cover."""
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    configure_styles(doc)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    set_paragraph_spacing(header, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_run_font(
        header.add_run("FINANCE RADAR  ·  INTERNSHIP PROJECT PROPOSAL  ·  V6.0"),
        size=8.5,
        bold=True,
        color=MUTED,
    )
    footer = section.footer.paragraphs[0]
    set_paragraph_spacing(footer, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_run_font(
        footer.add_run("北京林业大学理学院2026实训  ·  只读情报 / 无交易能力  ·  "),
        size=8.5,
        color=MUTED,
    )
    add_page_number(footer)


def add_display_para(
    doc: Document,
    text: str,
    *,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
    before: float = 0,
    after: float = 8,
    line: float = 1.0,
    size: float = 11,
    color=NAVY,
    bold: bool = False,
    italic: bool = False,
):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=before, after=after, line=line, align=align)
    set_run_font(
        paragraph.add_run(text),
        size=size,
        color=color,
        bold=bold,
        italic=italic,
    )
    return paragraph


def add_cover(doc: Document) -> None:
    add_display_para(
        doc,
        "北京林业大学理学院 2026 实训 · 自主高难度创新项目申报材料",
        before=30,
        after=20,
        size=10.5,
        bold=True,
        color=BLUE,
    )
    add_display_para(doc, "FINANCE RADAR", after=4, size=30, bold=True, color=NAVY)
    add_display_para(doc, "金融事件证据雷达", after=10, size=22, bold=True, color=NAVY)
    add_display_para(
        doc,
        "把“听说发生了什么”变成“哪一段原文证明了什么”",
        after=24,
        line=1.1,
        size=13,
        color=DARK_BLUE,
    )
    add_table(
        doc,
        ["材料版本", "成品基线", "正式部署", "项目边界"],
        [["V6.0", "2026-07-22", "AWS / HTTPS", "只读 · 无交易"]],
        [1700, 2100, 2600, 2960],
        font_size=9.3,
        center_cols=(0, 1, 2, 3),
    )
    add_callout(
        doc,
        "一句话说明",
        "系统自动收集监管公告、公司新闻和宏观信息，保存原文和时间线，找出支持或反驳关键说法的证据；证据不足时拒绝下结论，并把需要关注的事件交给人复核。",
        fill=LIGHT_BLUE,
    )
    add_callout(
        doc,
        "先说明它不是什么",
        "它不是荐股软件，不预测涨跌，不连接资金、持仓或下单接口。行情只用于事件发生后的只读观察，任何模型结论都不能自动变成交易动作。",
        fill=LIGHT_GOLD,
        accent=GOLD,
    )
    add_display_para(
        doc,
        "适合零基础读者的项目书 / 人读版",
        before=14,
        after=4,
        size=10,
        italic=True,
        color=MUTED,
    )
    add_display_para(
        doc,
        "团队：__________    指导教师：__________    日期：__________",
        after=0,
        size=9.5,
        color=MUTED,
    )
    doc.add_page_break()


def build() -> Path:
    doc = Document()
    configure_document(doc)
    bullet_id, decimal_id = add_num_defs(doc)
    props = doc.core_properties
    props.title = "Finance Radar 金融事件证据雷达 V6.0 实训项目书"
    props.subject = "北京林业大学理学院2026实训自主高难度创新项目申报材料"
    props.author = "Finance Radar Project Team"
    props.keywords = "Finance Radar, Evidence Agent, Web Terminal, Replay, Audit, No Trading"
    add_cover(doc)

    add_heading(doc, "执行摘要", 1)
    add_para(
        doc,
        "Finance Radar 是一套已经部署在 AWS 上持续运行的金融事件情报与证据系统。它从 22 个已登记来源中接收信息，把同一事件的不同消息、修订和官方原文组织成可追溯账本，再通过网页终端展示事件、证据、时间线、模型状态和系统健康。Telegram 只是提醒渠道，网页终端才是主操作界面。",
    )
    add_callout(
        doc,
        "立项判断",
        "本项目明显高于普通“多源新闻聚合 Agent”：它已经形成采集、核验、存证、回放、模型治理、网页展示、通知、备份恢复的完整工程链路。建议申报自主高难度创新项目；保守能力不低于标准级 95 分项目，但是否进入 100 分档必须由教师当面评审，项目书不预先承诺分数。",
        fill=LIGHT_GREEN,
        accent=GREEN,
    )

    add_heading(doc, "1. 不懂金融也能看懂：一个例子", 1)
    add_para(
        doc,
        "假设某个新闻账号突然发文：“A 公司被美国证券监管机构起诉”。普通聚合器往往只显示标题；Finance Radar 会继续追问：是谁发布的？监管机构原文在哪里？原文到底指控了什么？是否只是传闻？后来有没有官方更正？",
    )
    add_table(
        doc,
        ["步骤", "系统做什么", "人最终看到什么"],
        [
            ("1 发现", "从发现型新闻源或官方源收到消息", "一条待核验事件"),
            ("2 核对", "寻找 SEC 等一手来源并读取原文", "来源等级与原文链接"),
            ("3 存证", "保存精确段落、网页快照、时间和 SHA-256", "“哪段话证明哪件事”"),
            ("4 判断", "检查证据是否支持、反驳、冲突或不足", "核验 / 人工复核 / 弃权"),
            ("5 展示", "把事件、证据和过程放进网页工作台", "完整时间线与证据矩阵"),
            ("6 提醒", "满足严格条件后发送 Telegram 深链接", "点击即可回到网页证据页"),
        ],
        [1100, 4660, 3600],
        font_size=9,
        center_cols=(0,),
    )
    add_para(
        doc,
        "如果一家公司发布正面业绩，系统同样会保存，但不会把它强行解释成利空；如果消息互相矛盾，系统会显示冲突并要求人工复核。系统的价值不是“每次都给答案”，而是知道什么时候不能给答案。",
    )

    add_heading(doc, "2. 项目背景、目标与使用者", 1)
    add_heading(doc, "2.1 要解决的问题", 2)
    for item in [
        "金融信息分散在监管公告、公司新闻室、聚合平台和社交媒体中，来源质量差别很大。",
        "同一事件会经历传闻、官方确认、修订和撤回，单看一个标题容易误判。",
        "大模型可以总结文字，却可能编造引用、忽略冲突，不能单独承担事实核验。",
        "课堂答辩具有随机性：现场不一定刚好发生重大 SEC 事件，因此需要可重复回放。",
    ]:
        add_list(doc, item, bullet_id)
    add_heading(doc, "2.2 项目目标", 2)
    for item in [
        "建立多源、分级、可持续运行的事件采集链路；",
        "把新闻拆成可逐条验证的声明，并绑定精确证据；",
        "在证据不足或冲突时强制弃权，而不是让模型猜；",
        "提供专业网页终端、Telegram 提醒、历史回放和运维状态；",
        "留下测试、日志、Git、备份和恢复证据，满足软件工程实训要求。",
    ]:
        add_list(doc, item, decimal_id)
    add_heading(doc, "2.3 目标使用者", 2)
    add_para(
        doc,
        "主要使用者是需要快速理解重大公司、监管和宏观事件的研究者、学生和审核人员。普通读者可通过事件卡片理解“发生了什么”；专业人员可继续查看原文、引用、版本、模型路由和审计记录。",
    )

    add_heading(doc, "3. 系统由哪些部分组成", 1)
    add_callout(
        doc,
        "一条信息的完整旅程",
        "事件源 → 自动采集 → 去重与版本化 → 原文存证 → 证据关系判断 → 人工复核 → 网页/Telegram → 回放与备份。",
        fill=LIGHT_BLUE,
    )
    add_table(
        doc,
        ["层次", "主要组件", "通俗解释"],
        [
            ("入口", "SEC、CFTC、Fed、ECB、OpenNews 等", "新闻和公告从哪里来"),
            ("采集", "Python Worker / systemd", "每 5 分钟自动查看有没有新内容"),
            ("账本", "SQLite Ledger Schema 12", "保存原始消息、事件、版本和证据关系"),
            ("证据仓", "Operations Schema 4 + SHA-256 对象", "保存官方原文和不可变快照"),
            ("Agent", "确定性证据门 + 本地小模型摘要", "会总结，但不能越过证据规则"),
            ("接口", "FastAPI 只读 API", "把数据安全地提供给网页"),
            ("终端", "Evidence Terminal v2", "五个页面展示态势、证据、回放和运维"),
            ("恢复", "在线备份 + AES-256-GCM 异机备份", "服务器损坏后仍能恢复"),
        ],
        [1200, 3300, 4860],
        font_size=8.9,
    )
    add_heading(doc, "3.1 事件源与行情源必须分开", 2)
    add_table(
        doc,
        ["类别", "示例", "能证明什么 / 不能证明什么"],
        [
            ("P0 一手来源", "SEC、CFTC、FTC、FDIC、Fed、BLS、FDA、ECB、EIA", "可证明公告或申报内容；不能自动证明价格一定如何变化"),
            ("P1 发行人来源", "公司新闻室，如 NVIDIA", "可证明公司说过什么；不保证内容完全中立"),
            ("P2 发现来源", "OpenNews 等聚合发现源", "用于发现线索；不能直接升级为已核验事实"),
            ("行情观察", "Twelve Data、Binance 公共行情、IBKR 本机探针", "只记录事件后的价格窗口；不能充当事实证据或交易信号"),
        ],
        [1800, 3000, 4560],
        font_size=8.8,
    )

    add_heading(doc, "4. 使用者看到的五个页面", 1)
    add_table(
        doc,
        ["页面", "主要内容", "答辩时证明什么"],
        [
            ("态势室", "事件流、数量、来源健康、证据对象、Worker 和备份", "系统真实在线且在持续采集"),
            ("事件工作台", "事件详情、声明—证据矩阵、市场观察、冲突门", "判断来自哪条原文，而不是模型一句话"),
            ("回放实验室", "固定真实案例按时间推进", "没有现场新闻也能重复演示完整链路"),
            ("运行与模型", "服务、来源、模型盲测、备份恢复", "失败和限制也公开，不只展示成功"),
            ("盲标裁决", "双人互盲、第三人裁决、写入口关闭", "训练标签由人审内容产生，不由模型自证"),
        ],
        [1850, 4550, 2960],
        font_size=8.9,
    )
    add_callout(
        doc,
        "三种演示模式",
        "LIVE 证明外部连接和最新采集；RECENT_CAPTURE 展示真实历史；REPLAY 用固定案例证明证据如何逐步补齐。答辩成功不要求现场恰好出现重大事件。",
        fill=LIGHT_GREEN,
        accent=GREEN,
    )

    add_heading(doc, "5. AI 在项目里到底做什么", 1)
    add_para(
        doc,
        "本项目把 AI 放在核心流程中，但不把最终权力交给 AI。系统中的 AI 有两个职责：第一，把大量事件优先级排队，帮助人先看可能重要的下行风险；第二，在证据齐全时生成带引用的简短摘要。事件是否已经被证实、证据是否冲突、能否发送提醒，仍由确定性规则和人工审核控制。",
    )
    add_table(
        doc,
        ["能力", "允许", "禁止"],
        [
            ("风险路由模型", "RISK_REVIEW / NON_TARGET / ABSTAIN", "LONG、SHORT、目标价、收益预测"),
            ("本地摘要模型", "根据已经提供的证据生成 advisory summary", "编造引用、自动核验事实、覆盖人工结论"),
            ("证据 Agent", "拆声明、规划证据、提出支持/反驳关系", "证据不足时强行输出结论"),
        ],
        [2000, 3800, 3560],
        font_size=9,
    )
    add_heading(doc, "5.1 为什么模型优先训练负面事件", 2)
    add_para(
        doc,
        "负面重大事件通常具有损失不对称和更高的审核时效性，所以“优先发现需要人工复核的下行风险”是合理任务。但系统仍全极性采集正面、中性和混合事件；正面新闻一般进入 NON_TARGET 或 ABSTAIN，不会被解释成做空信号。",
    )
    add_heading(doc, "5.2 诚实展示失败比假装准确更重要", 2)
    add_para(
        doc,
        "现有小模型在 40 条 label-first 外部盲测中风险召回率为 100%，但正常新闻误报率达到 95%，因此门禁为 FAIL。系统没有隐藏结果，也没有拿这 40 条盲测反向调参；模型继续保持 SHADOW，只作为队列助手。这一失败反而证明项目具备真实模型治理，而不是只报一个漂亮准确率。",
    )

    add_heading(doc, "6. 当前已经完成到什么程度", 1)
    add_table(
        doc,
        ["项目", "2026-07-22 验收快照", "状态"],
        [
            ("正式部署", "AWS HTTPS；release 20260721T184054Z；API/Web/Worker/模型/备份服务运行", "已完成"),
            ("真实数据", "1,670 事件；2,399 证据边；1,595 原文证据对象；22 个来源", "持续增长"),
            ("数据完整性", "Ledger Schema 12、Operations Schema 4；两库 quick_check=ok", "已完成"),
            ("通知", "Telegram 完成安全切换，只发送未来符合条件的已核验事件", "已完成"),
            ("恢复", "加密 AWS 恢复包 779,640,430 字节；14,353 项完整恢复审计 PASS", "已完成"),
            ("测试", "364 tests + 17 subtests；GitHub CI 通过", "已完成"),
            ("课程人工证据", "24 条双人盲标仍为 0 人审；学生禁飞区、角色和计时练习尚待真实完成", "待学生"),
        ],
        [1800, 5900, 1660],
        font_size=8.5,
        center_cols=(2,),
        status_col=2,
    )
    add_callout(
        doc,
        "真实入口与版本",
        "网页终端：https://radar.18-208-34-152.sslip.io:8443/radar/ ；私有 GitHub 版本：v2026.07.22.1。服务器会继续自动保存新事件和原文证据。",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "7. 为什么符合实训要求", 1)
    add_para(
        doc,
        "学校要求从“会用 AI 做项目”升级为“能审查并驾驭 AI 产出”，强调完整工程链路、AI 深度融合、设计先行、细粒度 Git、测试驱动、Bug 注入、即兴修改和三处禁飞区。本项目与这些要求逐项对应。",
    )
    add_table(
        doc,
        ["学校要求", "本项目对应设计", "可验收证据"],
        [
            ("工程实践", "采集、API、网页、数据库、Worker、部署和恢复全流程", "AWS 成品、运行状态、恢复报告"),
            ("AI 代码审查", "模型失败门禁、捷径审计、确定性回退、禁止越权", "盲测报告、模型卡、测试"),
            ("工程规范", ".agent 目录、接口契约、Git、CI、自动测试、版本化发布", "仓库文件和提交历史"),
            ("团队协作", "PO/SM/QA 分工，所有成员开发，双人审核与第三人裁决", "角色表、Review、Sprint 记录"),
            ("AI 深度融合", "AI 参与证据拆解、摘要和风险路由，不是附加聊天框", "Agent trace、结构化输出、回放"),
            ("现场能力", "可注入状态、时间窗、证据冲突和来源失败类 Bug", "30 分钟修复与回归记录"),
        ],
        [2000, 4460, 2900],
        font_size=8.7,
    )
    add_callout(
        doc,
        "项目等级建议",
        "学校把“多源新闻聚合 Agent”列为标准级 95 分项目。本项目在其基础上增加精确证据、不可变原文、确定性回放、双人盲标、模型失败治理、AWS 运维和完整恢复，具备申报高难度创新级的理由；最终等级和评分上限仍以教师 Day 1 当面评审为准。",
        fill=LIGHT_GOLD,
        accent=GOLD,
    )

    add_heading(doc, "8. 12 天实施计划", 1)
    add_table(
        doc,
        ["时间", "目标", "主要工作与可见产物"],
        [
            ("Day 1", "方法与立项", "教师确认选题等级、团队角色、三处禁飞区；建立 Backlog"),
            ("Day 2-3", "理解现有系统", "画架构图、读 API/数据模型、运行测试、完成一次只读数据追踪"),
            ("Day 4-6 · Sprint 1", "核心手写模块", "禁飞区 FZ1/FZ2；先写失败测试再实现；完成第一次 Bug 注入"),
            ("Day 7-9 · Sprint 2", "闭环与联调", "禁飞区 FZ3；网页联调、双人 Review、即兴字段/规则修改"),
            ("Day 10-11 · Sprint 3", "质量与答辩", "性能、安全、回放、备份检查；每人三轮计时 Bug 练习"),
            ("Day 12", "最终验收", "三分钟演示、代码走查、现场 Bug/需求修改、反思报告"),
        ],
        [1850, 2600, 4910],
        font_size=8.8,
    )
    add_para(
        doc,
        "每个 Sprint 必须形成 Backlog、代码提交、测试、Review、缺陷记录和 AI 使用反思。已有工程成品是学习与集成基线，不得倒签为学生手写成果。",
    )

    add_heading(doc, "9. 团队角色与 AI 禁飞区", 1)
    add_heading(doc, "9.1 建议团队分工", 2)
    add_table(
        doc,
        ["角色", "主要职责", "仍需承担的开发任务"],
        [
            ("PO / 产品负责人", "确定优先级、用户故事和验收标准", "事件工作台或演示流程"),
            ("SM / Scrum Master", "站会、进度、障碍和 Sprint 复盘", "Worker、数据流或回放联调"),
            ("QA / 质量负责人", "测试设计、Bug 注入、代码 Review、AI 日志审核", "自动测试与质量仪表"),
            ("开发成员", "负责分配模块并参与 Review", "禁飞区、适配器、接口或 UI"),
        ],
        [1800, 3700, 3860],
        font_size=8.9,
    )
    add_heading(doc, "9.2 建议三处禁飞区（须教师批准）", 2)
    add_table(
        doc,
        ["禁飞区", "学生必须手写的核心", "验收方式"],
        [
            ("FZ1 事件身份与修订", "标准化、指纹、同一事件判断、修订链", "边界测试 + 逐行讲解"),
            ("FZ2 证据硬门", "支持/反驳/冲突/不足聚合，fail-closed 状态机", "失败测试先行 + 非法转换演示"),
            ("FZ3 确定性回放", "模拟时钟、证据逐步到达、输出可重复", "相同输入两次结果一致"),
        ],
        [2200, 4360, 2800],
        font_size=8.8,
    )
    add_callout(
        doc,
        "禁飞区规则",
        "AI 可以解释概念、提出测试思路和审查外围代码，但不得生成三处最终实现。学生需要保留设计草图、首次失败测试、细粒度提交、同伴 Review 和现场讲解记录。",
        fill=LIGHT_GOLD,
        accent=GOLD,
    )

    add_heading(doc, "10. Bug 注入与即兴修改准备", 1)
    add_table(
        doc,
        ["类型", "练习示例", "期待定位方法"],
        [
            ("基础型", "来源新鲜度单位写错、空值被显示为超大时长", "复现 → 最小测试 → 修正格式化"),
            ("边界型", "同一事件修订被错误当成新事件", "构造边界输入 → 检查指纹/版本链"),
            ("综合型", "冲突证据到达后仍发送提醒", "沿事件—证据—状态—outbox 全链路追踪"),
            ("即兴修改", "新增来源筛选、风险阈值或事件字段", "先改契约和测试，再改实现与 UI"),
        ],
        [1600, 4300, 3460],
        font_size=8.8,
    )
    add_para(
        doc,
        "统一答题顺序为：复现问题 → 找到证据 → 说明根因 → 做最小修复 → 跑回归测试 → 说明安全边界。每位成员至少完成 3 次 30 分钟 Bug 练习和 1 次即兴修改，留下真实时间与提交记录。",
    )

    add_heading(doc, "11. 验收标准与最终交付物", 1)
    add_heading(doc, "11.1 验收标准", 2)
    add_table(
        doc,
        ["方面", "通过标准", "当前状态"],
        [
            ("功能", "五页面、真实事件、回放、通知、备份可用", "工程已通过"),
            ("证据", "关键说法有精确引文；冲突/不足必须弃权或人审", "工程已通过"),
            ("安全", "无订单、仓位、余额、交易执行；秘密不进 Git", "工程已通过"),
            ("模型", "失败公开；未过门禁不得晋级；不得使用盲测调参", "保持 SHADOW"),
            ("工程", "测试、CI、细粒度提交、接口/数据/部署文档齐全", "基础已通过"),
            ("课程真实性", "禁飞区、角色、个人贡献、计时练习均为真实学生证据", "待实训完成"),
        ],
        [1700, 5000, 2660],
        font_size=8.8,
        center_cols=(2,),
        status_col=2,
    )
    add_heading(doc, "11.2 交付物", 2)
    _, delivery_numbers = add_num_defs(doc)
    for item in [
        "可运行代码、requirements、配置样例和启动脚本；",
        "AWS 网页终端、只读 API、Worker 与恢复说明；",
        "Schema 12/4 数据字典、接口契约和架构设计；",
        "真实 Replay、模型卡、盲测报告、测试与故障演练报告；",
        "人读项目书、AI 项目书、.agent 记忆目录和 AI 使用日志；",
        "团队角色、三处禁飞区、Git 贡献、Bug 注入和即兴修改证据；",
        "答辩 PPT、三分钟演示脚本和离线演示包。",
    ]:
        add_list(doc, item, delivery_numbers)

    add_heading(doc, "12. 三分钟答辩路线", 1)
    add_table(
        doc,
        ["时间", "动作", "要说清楚的结论"],
        [
            ("0:00-0:25", "打开态势室", "AWS 正在采集；事件、来源、证据和备份是真实状态"),
            ("0:25-1:20", "打开一个事件工作台", "指出精确原文、支持/反驳关系和证据不足门"),
            ("1:20-2:05", "运行固定 Replay", "事件从传闻到官方证据，现场无新闻也能重复"),
            ("2:05-2:35", "展示模型盲测 FAIL", "模型只做 Shadow 路由，失败不会被隐藏"),
            ("2:35-3:00", "展示恢复和 NO TRADING", "系统可恢复、无交易能力、边界可审计"),
        ],
        [1500, 3300, 4560],
        font_size=8.9,
    )

    add_heading(doc, "13. 风险、停止规则与结论", 1)
    add_heading(doc, "13.1 主要风险", 2)
    for item in [
        "外部来源可能限流、改版或暂时不可用：使用来源健康、退避、缓存和多源发现降低影响。",
        "模型可能跨来源失效：保持 Shadow、外部盲测、漂移门和人工复核。",
        "行情窗口可能错过：明确记录 MISSED_WINDOW，不用最新价格冒充历史价格。",
        "学生可能只会演示不会解释：通过禁飞区、代码走查、Bug 注入和个人提交验证理解。",
        "服务器可能更换：通过加密异机备份和完整恢复审计控制。",
    ]:
        add_list(doc, item, bullet_id)
    add_heading(doc, "13.2 停止扩展规则", 2)
    add_para(
        doc,
        "在教师审批、学生禁飞区、真实贡献和现场练习尚未完成前，不新增 Kubernetes、Redis、复杂前端框架或任何交易功能。继续堆技术名词不会提高课程成绩，补齐最短板才会。",
    )
    add_callout(
        doc,
        "最终结论",
        "Finance Radar 已具备“专业成品”的工程基础，也比普通新闻聚合 Agent 更有技术深度和可解释性。它有能力成为 95+ 作品；真正决定最终成绩的，将是学生能否理解并手写核心模块、留下真实过程、在现场修复问题，而不是服务器上已经有多少代码。",
        fill=LIGHT_GREEN,
        accent=GREEN,
    )

    add_heading(doc, "附录 A：术语表", 1)
    add_table(
        doc,
        ["术语", "最简单的解释"],
        [
            ("事件源", "告诉系统“可能发生了什么”的公告或新闻来源"),
            ("行情源", "提供股票、外汇或加密资产价格，只用于事后观察"),
            ("证据边", "一条声明与一段支持或反驳它的原文之间的连接"),
            ("不可变快照", "保存后按哈希校验、不能悄悄替换的原文副本"),
            ("Replay", "用同一批历史输入按时间重新播放，结果应可重复"),
            ("Shadow", "模型在旁边给建议，但没有最终决定权"),
            ("盲测", "模型训练时没见过、且先定答案再测试的数据"),
            ("硬门", "不满足条件就必须停止，不能让模型自由猜"),
        ],
        [2200, 7160],
        font_size=9,
    )

    add_heading(doc, "主要依据与参考资料", 1)
    add_reference(doc, 1, "北京林业大学理学院 2026 年实训动员会（课程文件）")
    add_reference(doc, 2, "北京林业大学理学院 2026 年项目列表（课程文件）")
    add_reference(doc, 3, "Finance Radar Technical Gap Closure 2026-07-22（项目验收记录）")
    add_reference(doc, 4, "SEC Developer Resources", "https://www.sec.gov/about/developer-resources")
    add_reference(doc, 5, "SQLite Online Backup API", "https://www.sqlite.org/backup.html")
    add_reference(doc, 6, "Model Cards for Model Reporting", "https://arxiv.org/abs/1810.03993")
    add_reference(doc, 7, "Datasheets for Datasets", "https://arxiv.org/abs/1803.09010")
    add_reference(doc, 8, "W3C PROV-O", "https://www.w3.org/TR/prov-o/")

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
