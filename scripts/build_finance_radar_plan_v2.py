#!/usr/bin/env python3
"""Build the human-readable Finance Radar V2.0 proposal."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "financial_event_radar_project_proposal_v2_1_human.docx"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1E293B"
MUTED = "667085"
LIGHT = "F4F6F9"
LIGHT_BLUE = "EAF2F8"
LIGHT_GOLD = "FFF7E6"
LIGHT_GREEN = "EDF7F0"
LIGHT_RED = "FDECEC"
WHITE = "FFFFFF"
GRID = "D7DEE8"


def set_run_font(run, *, size=11, bold=None, italic=None, color=INK, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, *, before=0, after=8, line=1.333, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.widow_control = True
    if align is not None:
        paragraph.alignment = align


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=GRID, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths_dxa):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {CONTENT_WIDTH_DXA}: {widths_dxa}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        existing = tbl_pr.find(qn(tag))
        if existing is not None:
            tbl_pr.remove(existing)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            set_cell_border(cell)


def add_num_defs(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abs = max(abstract_ids, default=0) + 1
    next_num = max(num_ids, default=0) + 1

    def create(kind, abs_id, num_id):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abs_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "280")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "290")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Calibri")
        fonts.set(qn("w:hAnsi"), "Calibri")
        r_pr.append(fonts)
        lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abs_id))
        num.append(abstract_ref)
        numbering.append(num)

    create("bullet", next_abs, next_num)
    create("decimal", next_abs + 1, next_num + 1)
    return next_num, next_num + 1


def set_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def add_para(doc, text="", *, bold_prefix=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             before=0, after=8, line=1.333, size=11, color=INK, keep=False):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=before, after=after, line=line, align=align)
    p.paragraph_format.keep_together = keep
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, size=size, bold=True, color=color)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest, size=size, color=color)
    else:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color)
    return p


def add_list(doc, text, num_id, *, size=11):
    p = doc.add_paragraph()
    set_num(p, num_id)
    set_paragraph_spacing(p, before=0, after=4, line=1.208, align=WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    set_run_font(run, size=size)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    p.add_run(text)
    return p


def add_table(doc, headers, rows, widths_dxa, *, header_fill=LIGHT, font_size=9.5,
              center_cols=()):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, header in enumerate(headers):
        cell = hdr.cells[idx]
        shade_cell(cell, header_fill)
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, before=0, after=0, line=1.05,
                              align=WD_ALIGN_PARAGRAPH.CENTER)
        run = p.add_run(header)
        set_run_font(run, size=9.5, bold=True, color=NAVY)
    for values in rows:
        row = table.add_row()
        for idx, value in enumerate(values):
            p = row.cells[idx].paragraphs[0]
            align = WD_ALIGN_PARAGRAPH.CENTER if idx in center_cols else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(p, before=0, after=0, line=1.08, align=align)
            run = p.add_run(str(value))
            set_run_font(run, size=font_size, color=INK)
    set_table_geometry(table, widths_dxa)
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=2, line=1.0)
    return table


def add_callout(doc, label, text, *, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=1, after=1, line=1.18, align=WD_ALIGN_PARAGRAPH.LEFT)
    r1 = p.add_run(f"{label}  ")
    set_run_font(r1, size=10.5, bold=True, color=accent)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=INK)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, after=4, line=1.0)
    return table


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, size=9, color=MUTED)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    set_paragraph_spacing(hp, before=0, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.LEFT)
    run = hp.add_run("FINANCE RADAR  ·  PROJECT PROPOSAL V2.0")
    set_run_font(run, size=8.5, bold=True, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    set_paragraph_spacing(fp, before=0, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    r = fp.add_run("内部研究与技术实施计划  ·  ")
    set_run_font(r, size=9, color=MUTED)
    add_page_number(fp)


def build():
    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    bullet_id, decimal_id = add_num_defs(doc)

    # Cover: proposal_centerpiece pattern with a restrained project-proposal treatment.
    add_para(doc, "PROJECT PROPOSAL  ·  V2.1", align=WD_ALIGN_PARAGRAPH.CENTER,
             before=36, after=16, line=1.0, size=10.5, color=BLUE)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=10, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("财经事件雷达 Agent")
    set_run_font(r, size=28, bold=True, color=NAVY)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=8, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("多源发现 · 快速候选 · 异步核验 · 逐资产影响 · 市场观察")
    set_run_font(r, size=14, color=DARK_BLUE)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=26, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("基于免费数据源与既有只读权限的个人财经事件研究系统")
    set_run_font(r, size=11, italic=True, color=MUTED)

    cover = doc.add_table(rows=2, cols=2)
    values = [
        ("版本", "V2.1", "编制日期", "2026年7月16日"),
        ("当前阶段", "M1主动官方源已实装 / 影子运行", "新增现金成本", "0元/月"),
    ]
    for row_idx, values_row in enumerate(values):
        for col_idx in range(2):
            label = values_row[col_idx * 2]
            value = values_row[col_idx * 2 + 1]
            cell = cover.rows[row_idx].cells[col_idx]
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, before=2, after=2, line=1.1, align=WD_ALIGN_PARAGRAPH.CENTER)
            r1 = p.add_run(f"{label}\n")
            set_run_font(r1, size=8.5, bold=True, color=MUTED)
            r2 = p.add_run(value)
            set_run_font(r2, size=10.5, bold=True, color=NAVY)
            shade_cell(cell, LIGHT)
    set_table_geometry(cover, [4680, 4680])

    add_para(doc, "", before=0, after=24, line=1.0)
    add_callout(
        doc,
        "核心决策",
        "Gate0不再等待所有候选API。项目正式进入M1，以OpenNews免费聚合发现、SEC/Fed/BLS官方核验、Binance/IBKR/Twelve Data行情观察和Telegram输出构成第一条垂直闭环。",
        fill=LIGHT_BLUE,
    )
    add_para(doc, "研究/监控用途，不构成投资建议；系统不具备也不调用自动交易权限。",
             align=WD_ALIGN_PARAGRAPH.CENTER, before=20, after=0, line=1.0,
             size=9.5, color=MUTED)
    doc.add_page_break()

    add_heading(doc, "执行摘要", 1)
    add_para(doc, "V2.0保留V1.0最有价值的部分：原文优先、来源分级、多维事件判断、逐步通知、行情口径诚实、审计与回放。同时吸收NewsLiquid/OpenNews成品架构中已经被公开证明有效的设计：原始新闻先到、AI结果异步补充；评分面向指定目标资产；实时分发与策略处理解耦。")
    add_para(doc, "本版本不追求复制一个交易终端，也不追逐单模型100毫秒的宣传数字。系统优先优化本地收到信息到发出候选提醒的端到端时延，同时让官方核验、事实版本、市场观察和更正机制在后台持续推进。")
    add_table(
        doc,
        ["维度", "V2.0决定", "结果"],
        [
            ("阶段", "Gate0结束，进入M1", "不再被可选密钥阻塞"),
            ("新闻", "OpenNews Free作为experimental P2", "立即补足聚合发现层"),
            ("处理", "快速路径与核验路径并行", "首次提示不等待完整分析"),
            ("评分", "事件级与逐资产级分离", "同一事件可对不同资产方向相反"),
            ("存储", "SQLite WAL + durable inbox/outbox", "零成本、可恢复、可升级"),
            ("安全", "行情只读、交易隔离", "不调用订单、持仓、资金接口"),
        ],
        [1440, 3960, 3960],
        center_cols=(0,),
    )

    add_heading(doc, "1. 为什么需要V2.0", 1)
    add_para(doc, "V1.0已经把项目边界、来源等级、事件生命周期、行情口径和风险控制写得相当完整，但逻辑架构仍然偏线性：采集后依次规范化、去重、判断、资产映射、行情增强和通知。对于实时新闻，这会让任何慢步骤拖住首次提醒。")
    add_para(doc, "同时，当前外部链路已经从“理论可用”进入“真实返回”：官方事件、聚合发现、加密行情、多资产行情和Telegram输出均至少有一条有效路径。因此下一版本必须从可行性论证转向可执行的流式系统设计。")
    add_callout(doc, "V2.0主张", "先落盘、先筛选、后核验；事件事实与资产影响分开；每次更新都有版本；每个外部依赖都可关闭或替换。", fill=LIGHT_GREEN, accent=DARK_BLUE)

    add_heading(doc, "2. 从NewsLiquid/OpenNews成品得到的结论", 1)
    add_heading(doc, "2.1 可确认的公开架构", 2)
    add_table(
        doc,
        ["层", "公开可确认事实", "对本项目的意义"],
        [
            ("采集", "85+来源，覆盖新闻、上币、链上、Meme、行情异常和预测", "聚合价值主要来自后端来源接入，不是MCP客户端"),
            ("接口", "REST查询 + WebSocket实时推送", "适配器需同时支持轮询和流式模式"),
            ("异步", "先推news.update，再推news.ai_update", "原始新闻不应等待模型处理"),
            ("评分", "每个目标资产单独给出score/signal/grade", "不能给事件保存一个全局利多/利空"),
            ("分发", "策略触发至少使用NATS按用户分发", "高吞吐系统需要消息总线；个人版可先用SQLite outbox"),
            ("模型", "专用窄任务评分器，非通用聊天模型", "优先严格接口和短输出，不追求长篇生成"),
        ],
        [1260, 4050, 4050],
    )
    add_heading(doc, "2.2 不应误读的部分", 2)
    for text in (
        "OpenNews公开了客户端与统一字段，但没有公开真实采集器、去重聚类、来源授权或模型训练细节。",
        "新闻来源名称不等于独立来源；Reuters原文、媒体转载和Twitter转发可能属于同一个origin。",
        "公开的140.1ms或270.3ms是单条新闻评分延迟，不是事件发布到用户收到提醒的完整延迟。",
        "公开基准使用200条新闻，评价即时解释而不是实际交易收益；项目不能拿该排行榜作为盈利证明。",
    ):
        add_list(doc, text, bullet_id)
    add_heading(doc, "2.3 采用与拒绝", 2)
    add_table(
        doc,
        ["采用", "拒绝"],
        [
            ("原始新闻先推、AI异步补充", "等待大模型完成才发首次提醒"),
            ("逐目标资产影响评分", "单一全局利多/利空"),
            ("流式分发、稳定ID、同事件更新", "每篇报道单独轰炸Telegram"),
            ("Provider分数保留为外部元数据", "把聚合器分数当事实或交易信号"),
            ("消息系统可替换", "个人规模提前引入Kafka/NATS复杂度"),
            ("研究与执行隔离", "复制NewsLiquid的交易执行集成"),
        ],
        [4680, 4680],
    )

    add_heading(doc, "3. 产品定位与不可逾越的边界", 1)
    add_para(doc, "财经事件雷达是个人研究基础设施，不是自动交易策略。它回答四个问题：发生了什么、是否可信、哪些资产可能受影响、市场是否出现可区分的反应。")
    for text in (
        "不调用任何订单、撤单、账户、持仓或资金接口。",
        "IBKR保持TWS Read-Only；新加坡服务器只运行固定的Binance公共行情请求。",
        "P2/P3只能产生候选；S/A已确认必须有P0/P1或满足独立交叉验证规则。",
        "不绕过付费墙、验证码、robots限制，不批量保存或转发新闻全文。",
        "市场反应属于观察结果，不属于事实确认，也不自动形成交易动作。",
    ):
        add_list(doc, text, bullet_id)

    add_heading(doc, "4. V2.0目标架构", 1)
    add_heading(doc, "4.1 三条数据路径", 2)
    add_table(
        doc,
        ["路径", "目标", "关键步骤", "用户结果"],
        [
            ("快速路径", "判断值不值得看", "落盘、轻量规范化、指纹、实体候选、快速评分", "FLASH · 待核验"),
            ("核验路径", "形成可审计事实", "事件聚类、来源独立性、官方原文、冲突检查、版本", "确认/争议/更正"),
            ("行情路径", "观察是否重新定价", "1/5/15/60分钟任务、收益、异常收益、波动与来源限定RVOL", "市场观察更新"),
        ],
        [1320, 1800, 3900, 2340],
    )
    add_heading(doc, "4.2 数据面与控制面", 2)
    add_table(
        doc,
        ["平面", "组件", "责任"],
        [
            ("数据面", "Adapters、Ingest、Fast Path、Verifier、Market、Alerts", "处理真实事件与行情"),
            ("控制面", "Source Registry、健康、配额、版本、重试、Replay、Labels", "决定如何运行、如何恢复、如何评估"),
        ],
        [1440, 3900, 4020],
    )
    add_heading(doc, "4.3 零成本消息架构", 2)
    add_para(doc, "M1采用进程内asyncio.Queue提升吞吐，SQLite WAL负责持久化。原始记录必须先提交，再创建pipeline job；Telegram采用durable outbox，使用(event_id, event_version, message_type)作为幂等键。进程重启后从PENDING/RETRY恢复，不重复发送已完成版本。")
    add_callout(doc, "升级条件", "只有当单机吞吐、延迟或多进程部署成为真实瓶颈时，才考虑NATS/Redis Streams/PostgreSQL；不以架构炫技增加M1复杂度。", fill=LIGHT_GOLD, accent="7A5A00")

    add_heading(doc, "5. 数据源方案与当前证据", 1)
    add_heading(doc, "5.1 Gate0真实状态", 2)
    add_table(
        doc,
        ["指标", "结果", "判断"],
        [
            ("PASS", "12", "真实返回并通过最低结构校验"),
            ("WARN", "3", "已有替代入口或需退避"),
            ("FAIL", "0", "不存在已调用但失败的核心能力"),
            ("BLOCKED", "5", "缺少可选密钥，不阻塞M1"),
        ],
        [1560, 1560, 6240],
        center_cols=(0, 1),
    )
    add_heading(doc, "5.2 当前可用来源", 2)
    add_table(
        doc,
        ["能力", "来源", "状态", "V2.0角色"],
        [
            ("宏观官方", "Fed RSS / BLS RSS+API", "PASS", "P0事实锚点"),
            ("公司监管", "SEC Submissions", "PASS", "P0事实锚点"),
            ("全球发现", "GDELT", "WARN 429", "低频P2备用"),
            ("聚合发现", "OpenNews Free", "人工实测PASS", "experimental P2"),
            ("加密行情", "Binance REST/WS", "PASS/WARN替代", "主行情适配器"),
            ("远程加密", "Singapore public relay", "PASS", "区域限制备用"),
            ("多资产行情", "Twelve Data", "PASS", "股票/ETF/FX/加密"),
            ("既有权限", "IBKR TWS Read-Only", "部分PASS", "FX/期货/授权资产"),
            ("通知", "Telegram Bot", "PASS", "唯一正式输出"),
        ],
        [1560, 2580, 1800, 3420],
        center_cols=(2,),
    )
    add_heading(doc, "5.3 OpenNews接入政策", 2)
    for text in (
        "免费接口立即接入，但标记experimental、P2、cached/periodic，不承诺实时。",
        "Provider score、grade、signal保存为provider_assessment，不覆盖本地判断。",
        "完整REST/WebSocket需要OPENNEWS_TOKEN；取得后必须完成48小时烘烤、断线补数、配额与条款核验。",
        "OpenNews可一键关闭；关闭后SEC/Fed/BLS等官方链路必须继续运行。",
        "默认只保存元数据、必要短摘要、哈希和链接。",
    ):
        add_list(doc, text, bullet_id)
    add_heading(doc, "5.4 来源独立性", 2)
    add_para(doc, "Source Registry新增provider_name、origin_source、independence_group、delivery_mode、storage_policy和latency_slo。重大事件的交叉验证按独立原始来源计算，不按文章数量、域名数量或聚合器返回条目数计算。")

    add_heading(doc, "6. 事件、资产与时间模型", 1)
    add_heading(doc, "6.1 事件事实与资产影响分离", 2)
    add_table(
        doc,
        ["对象", "保存内容", "禁止混入"],
        [
            ("Canonical Event", "发生了什么、类型、严重度、可信度、新颖度、事实状态", "全局利多/利空"),
            ("Asset Impact", "目标资产、关系、方向、影响、窗口、置信度、理由码", "事实是否真实"),
            ("Market Observation", "venue、data_scope、收益、异常收益、RVOL、波动", "因果结论"),
            ("Provider Assessment", "外部聚合器score/signal/summary", "本地最终判断"),
        ],
        [1800, 4920, 2640],
    )
    add_para(doc, "同一地缘事件可以对原油偏多、航空股偏空、黄金偏多、宽基指数偏空。逐资产评分发生在候选实体与资产解析之后，不能在资产映射前给事件一个统一方向。")
    add_heading(doc, "6.2 事实状态与任务状态分离", 2)
    add_table(
        doc,
        ["类型", "状态"],
        [
            ("事件事实状态", "NEW → VERIFYING → CONFIRMED / DISPUTED → MARKET_OBSERVED → MONITORING → RETRACTED / CLOSED"),
            ("后台任务状态", "PENDING → RUNNING → RETRY → DONE / DEAD"),
        ],
        [1800, 7560],
    )
    add_heading(doc, "6.3 全链路时间戳", 2)
    add_table(
        doc,
        ["字段", "含义"],
        [
            ("source_published_at", "来源标明的发布时间"),
            ("provider_observed_at", "聚合器首次观察；不可得时为空"),
            ("local_received_at", "本系统收到响应或WS消息"),
            ("persisted_at", "原始记录提交SQLite"),
            ("normalized_at / triaged_at", "轻量处理与快速评分完成"),
            ("confirmed_at", "达到事实确认条件"),
            ("alert_sent_at", "Telegram创建/更新成功"),
            ("market_observed_at", "指定行情观察窗完成"),
        ],
        [2880, 6480],
    )

    add_heading(doc, "7. 去重、聚类与版本", 1)
    add_para(doc, "去重不能只靠URL或语义向量。V2.0采用五层策略：")
    for text in (
        "URL规范化和外部ID。",
        "内容SHA-256精确去重。",
        "规范化标题与SimHash/MinHash近重。",
        "主体 + 动作 + 对象 + 关键数字 + 时间桶组成事件指纹。",
        "语义聚类只用于候选合并，不覆盖硬证据和来源独立性判断。",
    ):
        add_list(doc, text, decimal_id)
    add_para(doc, "同一来源编辑写入source_revisions；出现新金额、新法律阶段、时间表或官方回应时生成event_version。删除、撤回和官方否认不能静默消失，必须更新同一Telegram事件线程。")

    add_heading(doc, "8. 模型策略", 1)
    add_heading(doc, "8.1 三级级联", 2)
    add_table(
        doc,
        ["阶段", "预算", "任务", "失败处理"],
        [
            ("Stage 0规则", "P95 ≤ 50ms", "来源等级、词典、否定词、数字、哈希、垃圾过滤", "继续规则降级"),
            ("Stage 1快速评分", "P95 ≤ 500ms；P99 ≤ 1s", "事件候选、严重度候选、逐资产影响", "ABSTAIN/待复核"),
            ("Stage 2异步核验", "不阻塞首次提示", "官方检索、事实抽取、冲突、法律阶段、说明", "人工队列/延后更新"),
        ],
        [1620, 1980, 3900, 1860],
    )
    add_heading(doc, "8.2 逐资产输出契约", 2)
    add_table(
        doc,
        ["字段", "要求"],
        [
            ("target_asset", "明确的股票、ETF、期货、FX或加密资产"),
            ("direction", "LONG / SHORT / NEUTRAL / ABSTAIN"),
            ("impact", "0-100短期直接影响强度"),
            ("horizon", "例如15m-4h；不得省略"),
            ("confidence", "0-1；低于阈值不主动扩散资产"),
            ("reason_codes", "稳定、可回测的理由码"),
            ("assessment_source", "provider/local_rule/local_model/human"),
            ("model_version", "每次结果可追溯"),
        ],
        [2520, 6840],
    )
    add_callout(doc, "训练策略", "M1不训练自己的NewsLiquid。先建立规则基线、冻结样本和人工标签；OpenNews分数可作弱特征，绝不能作真值。", fill=LIGHT_GOLD, accent="7A5A00")

    add_heading(doc, "9. 行情观察与Telegram产品体验", 1)
    add_heading(doc, "9.1 市场观察", 2)
    add_para(doc, "行情路径使用Binance、Twelve Data和IBKR既有只读权限。每个指标必须携带venue、data_scope、data_as_of和stale标记。默认在1、5、15、60分钟及收盘后更新，休市时调度到下一可交易窗口。")
    add_table(
        doc,
        ["指标", "输出", "边界"],
        [
            ("收益", "1/5/15/60分钟", "标明交易所、交易对、复权和休市"),
            ("异常收益", "资产收益 - 基准收益", "基准缺失不计算，不解释为因果"),
            ("来源限定RVOL", "BINANCE_RVOL / IEX_RVOL", "字段名必须包含口径"),
            ("实现波动率", "数值 + 历史分位", "避免跨资产直接比较绝对值"),
            ("衍生品", "funding / OI / liquidation", "只代表指定交易所和合约"),
        ],
        [1800, 2880, 4680],
    )
    add_heading(doc, "9.2 Telegram事件线程", 2)
    for text in (
        "FLASH：快速候选，必须显示待核验。",
        "VERIFIED UPDATE：官方或独立可靠来源确认，编辑同一消息。",
        "MARKET UPDATE：按观察窗补充行情，不作为事实证明。",
        "CORRECTION：来源冲突、撤回或更正必须主动显示。",
        "DIGEST：B/C和非紧急更新进入摘要，避免告警疲劳。",
        "HEALTH：显示来源、配额、队列、失败、断线和行情新鲜度。",
    ):
        add_list(doc, text, bullet_id)

    add_heading(doc, "10. 数据、恢复与安全设计", 1)
    add_heading(doc, "10.1 核心数据表", 2)
    add_table(
        doc,
        ["域", "主要表"],
        [
            ("来源与原始", "sources、raw_observations、source_revisions"),
            ("事件与证据", "canonical_events、event_versions、event_observations、event_facts、event_conflicts"),
            ("实体与影响", "entities、assets、relations、asset_impacts"),
            ("行情", "market_jobs、market_snapshots"),
            ("任务与通知", "pipeline_jobs、dead_letters、alert_outbox、alerts"),
            ("模型与评估", "model_runs、labels、replay_runs"),
        ],
        [2160, 7200],
    )
    add_heading(doc, "10.2 恢复机制", 2)
    for text in (
        "WebSocket自动重连，并记录断线起止和最后游标。",
        "支持REST backfill或last_seen补数。",
        "RSS使用ETag和Last-Modified增量读取。",
        "任务指数退避，超过阈值进入dead letter。",
        "Outbox确保重启后不重复发送相同事件版本。",
    ):
        add_list(doc, text, bullet_id)
    add_heading(doc, "10.3 进程与权限隔离", 2)
    add_table(
        doc,
        ["进程", "允许", "禁止"],
        [
            ("Collector", "外部只读采集、来源Token", "账户、订单、密钥日志"),
            ("Market Reader", "公共行情、IBKR只读", "持仓、执行、账户、订单"),
            ("Event Engine", "本地数据库与规则/模型", "交易接口"),
            ("Notifier", "Telegram Bot API", "个人账号写操作、交易"),
        ],
        [1800, 3600, 3960],
    )

    add_heading(doc, "11. 实施路线", 1)
    add_table(
        doc,
        ["阶段", "主题", "主要交付", "退出条件"],
        [
            ("Gate0", "外部链路", "12 PASS、3 WARN、0 FAIL；OpenNews人工PASS", "完成"),
            ("M1", "最小流式闭环", "OpenNews/P0适配器、raw、inbox/outbox、Telegram", "48小时；重启不重发；可关闭OpenNews"),
            ("M2", "事件引擎", "去重聚类、版本、来源独立性、核验与更正", "冻结集首轮评估"),
            ("M3", "资产与行情", "Symbol Master、asset_impacts、行情窗口", "固定夹具100%一致"),
            ("M4", "模型与回放", "三级级联、模拟时钟、错误样本导出", "优于规则基线"),
            ("M5", "影子运行", "至少两周，不交易", "达到质量阈值或继续迭代"),
        ],
        [1260, 1800, 3780, 2520],
        center_cols=(0,),
    )
    add_heading(doc, "11.1 M1第一条垂直链路", 2)
    for text in (
        "接入OpenNews Free，同时保留SEC/Fed/BLS官方适配器。",
        "建立sources、raw_observations、pipeline_jobs、alert_outbox迁移。",
        "先落盘，再执行轻量规范化、事件指纹和资产候选。",
        "生成FLASH或确认事件卡，通过outbox发送/更新Telegram。",
        "连续运行48小时，记录各阶段时间戳、失败、重复率和数据新鲜度。",
    ):
        add_list(doc, text, decimal_id)

    add_heading(doc, "12. 验收指标", 1)
    add_table(
        doc,
        ["指标", "目标", "说明"],
        [
            ("本地持久化P95", "≤100ms", "收到数据到SQLite提交"),
            ("快速处理P95", "≤2s", "收到数据到triaged_at"),
            ("WS源本地提醒P95", "≤5s", "不含来源/聚合器自身延迟"),
            ("P0轮询捕获P95", "≤120s", "条目可获取到first_seen"),
            ("硬事件召回", "≥85%", "冻结hard/weak/rejected集"),
            ("S/A精确率", "≥90%", "人工复核"),
            ("错误已确认S率", "≤2%", "候选与已确认分开统计"),
            ("重复压缩率", "≥85%", "同事件重复报道"),
            ("聚类纯度", "≥95%", "冻结测试集"),
            ("PRIMARY资产准确率", "≥95%", "供应链关系另计"),
            ("确定性指标正确率", "100%", "固定行情夹具"),
            ("主动告警可追溯率", "100%", "来源、时间、口径、版本"),
        ],
        [2880, 1680, 4800],
        center_cols=(1,),
    )

    add_heading(doc, "13. 风险与应对", 1)
    add_table(
        doc,
        ["风险", "影响", "应对"],
        [
            ("聚合器改价/下线", "发现覆盖下降", "P2可关闭；P0独立运行；多适配器"),
            ("来源名称/授权不透明", "版权和确认错误", "存短摘要和链接；追溯origin；不当独立确认"),
            ("Provider分数漂移", "误报或方向偏差", "provider/local分离；冻结集监控"),
            ("WebSocket断线", "消息缺口", "游标、断线窗口、REST补数"),
            ("模型幻觉/法律混淆", "高风险错误", "结构化schema、证据句、硬规则、人工复核"),
            ("主体/资产歧义", "错误扩散", "CIK/代码/地址优先；ABSTAIN"),
            ("免费行情口径不足", "RVOL误导", "venue/data_scope强制字段；缺失N/A"),
            ("未来信息泄漏", "回测虚高", "事件版本、模拟时钟、t0前可见性"),
            ("告警疲劳", "用户忽视重要信息", "S/A主动；B/C摘要；同事件更新"),
        ],
        [2700, 2340, 4320],
    )

    add_heading(doc, "14. 交付物", 1)
    add_table(
        doc,
        ["编号", "交付物", "最低内容"],
        [
            ("D1", "Source Registry", "来源等级、能力、独立组、条款、配额、健康、开关"),
            ("D2", "采集适配器", "OpenNews Free、SEC、Fed、BLS；行情适配器复用现有探针"),
            ("D3", "SQLite迁移", "raw、事件版本、影响、任务、outbox、行情、模型、标签"),
            ("D4", "双速事件引擎", "Fast Path、Verification Path、修订与撤回"),
            ("D5", "逐资产与行情引擎", "asset_impacts、窗口任务、venue-scoped指标"),
            ("D6", "Telegram Bot", "FLASH、确认、市场、更正、摘要、健康"),
            ("D7", "Replay Evaluator", "模拟时钟、冻结集、错误导出、版本对比"),
            ("D8", "运行与安全文档", "安装、备份、恢复、断线、降级、只读边界"),
        ],
        [900, 2520, 5940],
        center_cols=(0,),
    )
    add_heading(doc, "14.1 当前下一任务", 2)
    add_callout(doc, "M1下一任务", "当前140个历史复核线程已经全部闭环。下一步不是继续堆通用API，而是生成下一批家族均衡、可审计候选，并把268条拒绝控制前置为发现端语义过滤和回归样本；实时端继续静默等待CBIO、Q32和Obsidian的真实交割文件。", fill=LIGHT_GREEN, accent=DARK_BLUE)

    add_heading(doc, "14.2 2026年7月16日实装验收", 2)
    add_para(doc, "主动官方事件层已经实跑，不再只是计划。Fed RSS、SEC EDGAR最新申报和BLS关键经济序列均进入统一账本；P0只提高来源可信度，不自动确认事件严重度。")
    add_table(
        doc,
        ["项目", "真实状态", "安全约束"],
        [
            ("Fed", "20条原始公告；按货币政策、监管、执法分类", "Other Announcements默认过滤"),
            ("SEC", "46条事件相关申报；按Form和Item路由", "正文与EX-99进入审核包"),
            ("BLS", "4个发布组；CPI、PPI、就业、JOLTS", "90分钟限频；修订可追溯"),
            ("统一账本", "Schema 12；2753条原始观察；710个来源版本", "不可变原始层和独立游标"),
            ("候选层", "68个candidate；历史triage为0", "候选不能自动升级"),
            ("SEC正文", "48条实时富化；账本证据1894条", "原文分类不自动核验"),
            ("审核分诊", "实时待审3；历史累计复核614", "排序分数不是严重度"),
            ("验证", "223个测试通过；19项审计违规均为0", "安全计数必须为0"),
        ],
        [1800, 4200, 3360],
    )
    add_callout(doc, "当前结论", "最需要做的仍不是堆来源，而是把实体、证券、事件日期、融资链、司法阶段和旧普通股终局对齐。尾部35个线程已全部关闭；下一阶段以新批次覆盖价值、发现端误报拦截率和可训练负样本质量为主门槛。", fill=LIGHT_BLUE, accent=DARK_BLUE)

    add_heading(doc, "14.3 主动历史研究与最终状态跟踪", 2)
    add_para(doc, "历史路径已经从一次性抽样升级为可持续研究循环：150条家族均衡队列每次推进25条，游标持久化、证据幂等合并，D:\\short始终只读。事后收益只做审计，不能进入发现排序或等级判定。")
    add_table(
        doc,
        ["质量指标", "当前值", "含义"],
        [
            ("历史队列扫描", "150 / 150", "已累计2511份SEC候选、1622条证据段"),
            ("关键词证据覆盖", "122 / 140线程", "87.1%线程已有可定位的一手正文"),
            ("本轮人工裁决", "140 / 140线程", "当前triage为0；累计复核614条"),
            ("累计裁决结果", "346 verified；268 rejected", "负样本与硬事件都保留证据链"),
            ("统一账本", "787规范事件", "451 verified；268 rejected；68 candidate"),
        ],
        [2400, 1920, 5040],
    )
    add_callout(doc, "价格信号只负责发现", "最后16个价格线程全部拒绝作为事件标签。HAO、EZGO和LNKS是既有事件的重复市场或复权后果；OST、MGN和ELPW的真实融资，以及Town Sports和Ability的司法重组，均按官方日期另建事件。没有任何价格跌幅直接决定等级。", fill=LIGHT_GOLD, accent="7A5A00")
    add_callout(doc, "来源错配已经清零", "普通SPAC信托赎回、目标公司事件传播到SPAC、后续OTC代码倒灌、Form 25后果和未来破产回填均已成为拒绝控制。LILM、VEV、GENE等真实困境事件按法律日期另建；AAMCF、CPTAF、DXIEF等后续代码不得覆盖AAMC、CAPT、DXI.H事件时点身份。", fill=LIGHT_BLUE, accent=DARK_BLUE)
    add_callout(doc, "反向拆股不能机械定级", "MAXN、CANO和CHUC的比例调整保留为B；ASTI、VLCN、SBFM、AWIN、KAL、NUWE和DBGI只有在一手文件证明已实现发行、预融资/重置权证、可转债违约或授权股本扩张链时才进入A。MAXNQ、EMPD、CANOQ、KALRQ是后续身份代理，已恢复为事件时点MAXN、VLCN、CANO、KAL。没有旧股终局损失证据，不得升A++或S。", fill=LIGHT_GOLD, accent="7A5A00")
    add_callout(doc, "退市原因先于等级", "EM和MRCC是有现金或股票对价的并购退市，不能当作股权死亡；SEAC、BNSO是going-dark；ABB、CAJ、CEA、ZNH保留境外主要上市；DTEA与PTNR分别转至TSXV和TASE。后续OTC代码只作别名，不能覆盖事件时点身份。", fill=LIGHT_GREEN, accent=DARK_BLUE)
    add_callout(doc, "双窗口规则", "同期窗口（-10至+45天）确认事件事实和初始状态；对拟议重组、Chapter 11、退市申诉再运行0至+180天最终状态跟踪。第二窗口只解决法律和旧股结局，不用于事后收益定级。", fill=LIGHT_GREEN, accent=DARK_BLUE)
    add_para(doc, "NINEQ说明了为什么必须这样做：初始8-K只写计划拟注销旧股，103天后的10-Q才证明计划已生效、旧普通股无对价注销、新普通股发给票据持有人。因此只有后续一手证据允许将其裁为S。")
    add_para(doc, "审核队列还增加两项反误报规则：同一证券、日期和事件家族的多个检测器合并为一个人工线程；Chapter 7单独识别，不能因法律引用中的“title 11”误分为Chapter 11。WOLF2进一步证明旧股被取消但获得新股恢复时不能升S。")
    add_para(doc, "退市研究新增-60至+30天原因窗口并始终比较主申报与EX-99正文：SPAC单位终止但继任证券继续上市应作为误报对照；双重上市整合或转OTC通常只是B；只有在低流动性、资本市场通道明显收缩且替代市场不确定时才进入A边界。")
    add_para(doc, "跨事件家族也要去重：同一稳定证券、同一日期的破产与破产驱动退市共用一个事件链。主破产是primary，退市保留为consequence，后者不再重复计入S/A++硬标签。")
    add_para(doc, "本轮对Enviva、Spirit、Vertex Energy、Edgio和Tupperware建立了双时点标签：申请日只到A++，只有计划正式生效且旧普通股被无分配注销时才在生效日记录S。Sharadar ACTIONS日期不能直接当成法律事件日期。")
    add_para(doc, "三类新增硬负样本已进入回归测试：披露文件中的假设性Chapter 7清算分析、已存在Chapter 11案件的重复风险表述、以及破产引发的交易所摘牌通知。它们分别是对照文本、旧状态和事件后果，不是新的主破产事件。")

    add_heading(doc, "15. 主要研究来源", 1)
    sources = [
        ("NewsLiquid产品与模型", "https://newsliquid.com/"),
        ("NewsLiquid 2.0模型说明", "https://app.newsliquid.com/blog/newsliquid-2-0-flash.html"),
        ("FinTech News Impact Benchmark v1", "https://app.newsliquid.com/blog/leaderboard.html"),
        ("OpenNews MCP官方仓库", "https://github.com/6551team/opennews-mcp"),
        ("OpenNews Free API说明", "https://raw.githubusercontent.com/6551team/opennews-mcp/main/openclaw-skill/opennews/SKILL.md"),
        ("SEC Developer Resources", "https://www.sec.gov/edgar/sec-api-documentation"),
        ("Federal Reserve RSS", "https://www.federalreserve.gov/feeds/feeds.htm"),
        ("BLS Public Data API", "https://www.bls.gov/developers/"),
        ("Binance Spot API Docs", "https://github.com/binance/binance-spot-api-docs"),
        ("IBKR TWS API", "https://interactivebrokers.github.io/tws-api/"),
    ]
    add_table(doc, ["来源", "URL"], sources, [3000, 6360], font_size=8.8)

    add_heading(doc, "立项结论", 1)
    add_para(doc, "财经事件雷达V2.0技术上可行，且已经具备进入实现阶段所需的外部能力。最适合本项目的路线不是复制NewsLiquid的交易终端，而是借用其异步流式、逐资产感知和实时分发思路，再叠加本项目更严格的官方核验、事件版本、行情口径和交易隔离。")
    add_para(doc, "最终架构确定为：多源发现 → 原始落盘 → 快速候选 → 异步核验 → 逐资产影响 → 市场观察 → 同事件更新。Gate0完成，M1立即开始。", bold_prefix="最终架构确定为：")
    add_para(doc, "— 文档结束 —", align=WD_ALIGN_PARAGRAPH.CENTER, before=24, after=0,
             line=1.0, size=9.5, color=MUTED)

    # Metadata and document settings.
    doc.core_properties.title = "财经事件雷达 Agent 项目计划书 V2.1"
    doc.core_properties.subject = "基于免费数据源的双速财经事件发现、核验与市场观察系统"
    doc.core_properties.author = "Finance Radar Project"
    doc.core_properties.keywords = "finance radar, event intelligence, OpenNews, Telegram, read-only market data"
    doc.core_properties.comments = "Updated with unified research-quality gates, continuous Sharadar/SEC review cycles, and 180-day final-state follow-up."
    doc.settings.element.append(OxmlElement("w:updateFields"))

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
